from flask import request, jsonify, send_file
import requests
import re
import os
import ast
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from sonarqube_security_standards import SecurityStandards, SQCategory, VulnerabilityProbability
from datetime import datetime
import json
import tempfile

# Based on SonarQube's SecurityStandards.java
COMMAND_INJECTION_VULNERABILITY = ("command-injection", VulnerabilityProbability.HIGH)
# Maps to CWE-77, CWE-78, CWE-88, CWE-214
# Maps to OWASP A03:2021-Injection

class CommandInjectionVulnerability:
    def __init__(self, line_number, vulnerability_type, description, severity, code_snippet, remediation, confidence, file_path=None):
        self.line_number = line_number
        self.vulnerability_type = vulnerability_type
        self.description = description
        self.severity = severity
        self.code_snippet = code_snippet
        self.remediation = remediation
        self.confidence = confidence
        self.file_path = file_path or 'unknown'
        
    def to_dict(self):
        return {
            'line_number': self.line_number,
            'vulnerability_type': self.vulnerability_type,
            'description': self.description,
            'severity': self.severity,
            'code_snippet': self.code_snippet,
            'remediation': self.remediation,
            'confidence': self.confidence,
            'file_path': self.file_path,
            'cwe_references': ["77", "78", "88", "214"],
            'owasp_references': ["A03:2021-Injection"],
            'rule_key': 'python:S2076'  # Command injection rule key similar to SonarQube
        }

class CommandInjectionDetector:
    def __init__(self):
        self.vulnerabilities = []
        
    def scan_file(self, filename):
        """Scan a Python file for command injection vulnerabilities"""
        self.vulnerabilities = []
        
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                code = f.read()
        except UnicodeDecodeError:
            with open(filename, 'r', encoding='latin-1') as f:
                code = f.read()
        
        # Scan for command injection patterns
        self._scan_for_patterns(code, filename)
        self._scan_with_ast(code, filename)
        
        # Remove duplicates
        self.vulnerabilities = self._deduplicate_vulnerabilities(self.vulnerabilities)
        
        return self.vulnerabilities
    
    def _deduplicate_vulnerabilities(self, vulnerabilities):
        """Remove duplicate vulnerabilities based on line number and similarity"""
        if not vulnerabilities:
            return vulnerabilities
        
        # Group vulnerabilities by line number
        by_line = {}
        for vuln in vulnerabilities:
            line_key = vuln.line_number
            if line_key not in by_line:
                by_line[line_key] = []
            by_line[line_key].append(vuln)
        
        # For each line, keep only the best vulnerability
        deduplicated = []
        for line_num, line_vulns in by_line.items():
            if len(line_vulns) == 1:
                deduplicated.extend(line_vulns)
            else:
                # Multiple vulnerabilities on the same line - keep the best one
                # Sort by confidence (highest first), then by description length (longest first)
                best_vuln = max(line_vulns, key=lambda v: (v.confidence, len(v.description)))
                deduplicated.append(best_vuln)
        
        return deduplicated
    
    def _scan_for_patterns(self, code, filename):
        """Scan for command injection vulnerability patterns using regex"""
        lines = code.split('\n')
        
        # Command injection vulnerability patterns
        patterns = [
            # os.system() with user input
            {
                'pattern': r'(os\.system\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'os.system() with concatenated user input can lead to command injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use subprocess.run() with shell=False and a list of arguments instead of os.system()'
            },
            {
                'pattern': r'(os\.system\s*\(\s*(?![\'"])[^)]*\w+[^)]*\))',
                'description': 'os.system() with potentially unsafe user input',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use subprocess.run() with shell=False and validate all inputs'
            },
            # subprocess.call() with shell=True
            {
                'pattern': r'(subprocess\.call\s*\([^)]*shell\s*=\s*True[^)]*\))',
                'description': 'subprocess.call() with shell=True can lead to command injection',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use shell=False and pass commands as a list of arguments'
            },
            {
                'pattern': r'(subprocess\.run\s*\([^)]*shell\s*=\s*True[^)]*\))',
                'description': 'subprocess.run() with shell=True can lead to command injection',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use shell=False and pass commands as a list of arguments'
            },
            {
                'pattern': r'(subprocess\.Popen\s*\([^)]*shell\s*=\s*True[^)]*\))',
                'description': 'subprocess.Popen() with shell=True can lead to command injection',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use shell=False and pass commands as a list of arguments'
            },
            # os.popen() patterns
            {
                'pattern': r'(os\.popen\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'os.popen() with concatenated user input can lead to command injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use subprocess.run() with shell=False and proper input validation'
            },
            {
                'pattern': r'(os\.popen\s*\(\s*(?![\'"])[^)]*\w+[^)]*\))',
                'description': 'os.popen() with potentially unsafe user input',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use subprocess.run() with shell=False and validate all inputs'
            },
            # os.spawn* family functions
            {
                'pattern': r'(os\.spawn[lv]p?\s*\([^)]*\))',
                'description': 'os.spawn* functions can lead to command injection if not properly validated',
                'severity': 'high',
                'confidence': 0.7,
                'remediation': 'Use subprocess.run() with shell=False and validate all inputs'
            },
            # eval() with command-like patterns
            {
                'pattern': r'(eval\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'eval() with user input can lead to code/command injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Avoid eval() completely or use safe alternatives'
            },
            {
                'pattern': r'(eval\s*\(\s*[^)]*\+\s*\w+[^)]*\))',
                'description': 'eval() with string concatenation and user input can lead to code/command injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Avoid eval() completely or use safe alternatives'
            },
            {
                'pattern': r'(eval\s*\(\s*\w+[^)]*\))',
                'description': 'eval() with potentially unsafe user input',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Avoid eval() completely or use safe alternatives'
            },
            {
                'pattern': r'(exec\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'exec() with user input can lead to code/command injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Avoid exec() completely or use safe alternatives'
            },
            {
                'pattern': r'(exec\s*\(\s*[^)]*\+\s*\w+[^)]*\))',
                'description': 'exec() with string concatenation and user input can lead to code/command injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Avoid exec() completely or use safe alternatives'
            },
            {
                'pattern': r'(exec\s*\(\s*\w+[^)]*\))',
                'description': 'exec() with potentially unsafe user input',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Avoid exec() completely or use safe alternatives'
            },
            # Shell command patterns in strings
            {
                'pattern': r'([\'"][^\'\"]*(?:sudo|su|bash|sh|cmd|powershell|python|perl|ruby|php)[^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Shell command construction with user input detected',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use subprocess with shell=False and validate inputs'
            },
            # Django/Flask command execution patterns
            {
                'pattern': r'(management\.call_command\s*\(\s*[\'"][^\'\"]*[\'"]\s*,\s*\w+)',
                'description': 'Django management command with user input can lead to command injection',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Validate and sanitize all user inputs before passing to management commands'
            },
            # File operation patterns that could be abused
            {
                'pattern': r'(os\.remove\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'File deletion with user input can lead to path traversal or command injection',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Validate file paths and use os.path.join() for safe path construction'
            },
            {
                'pattern': r'(os\.rmdir\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Directory removal with user input can lead to path traversal attacks',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Validate directory paths and restrict to allowed directories'
            },
            # Template injection patterns that could lead to command execution
            {
                'pattern': r'(Template\s*\(\s*[\'"][^\'\"]*\$[^\'\"]*[\'"]\s*\)\.substitute\s*\([^)]*\))',
                'description': 'Template substitution with user input can lead to injection attacks',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Use safe template engines and validate all template inputs'
            },
            {
                'pattern': r'(Template\s*\([^)]*\)\.substitute\s*\([^)]*\))',
                'description': 'Template substitution with potentially unsafe user input',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Use safe template engines and validate all template inputs'
            },
            {
                'pattern': r'(\.substitute\s*\([^)]*\))',
                'description': 'Template substitution with potentially unsafe user input',
                'severity': 'medium',
                'confidence': 0.6,
                'remediation': 'Use safe template engines and validate all template inputs'
            },
            # Configuration file execution patterns
            {
                'pattern': r'(configparser\.[^.]*\.read\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Configuration file reading with user input can lead to file inclusion attacks',
                'severity': 'medium',
                'confidence': 0.6,
                'remediation': 'Validate configuration file paths and restrict to allowed directories'
            },
            # Import/module loading patterns
            {
                'pattern': r'(__import__\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Dynamic import with user input can lead to code injection',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Validate module names against a whitelist before importing'
            },
            {
                'pattern': r'(__import__\s*\(\s*\w+[^)]*\))',
                'description': 'Dynamic import with potentially unsafe user input',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Validate module names against a whitelist before importing'
            },
            {
                'pattern': r'(importlib\.import_module\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Dynamic module import with user input can lead to code injection',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Validate module names against a whitelist before importing'
            },
            {
                'pattern': r'(importlib\.import_module\s*\(\s*\w+[^)]*\))',
                'description': 'Dynamic module import with potentially unsafe user input',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Validate module names against a whitelist before importing'
            },
            # JavaScript/Node.js patterns (for polyglot projects)
            {
                'pattern': r'(child_process\.exec\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'child_process.exec() with user input can lead to command injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use child_process.spawn() or child_process.execFile() with validated inputs'
            },
            {
                'pattern': r'(require\s*\(\s*[\'"]child_process[\'\"]\s*\)\.exec)',
                'description': 'Use of child_process.exec() detected - potential command injection risk',
                'severity': 'medium',
                'confidence': 0.6,
                'remediation': 'Ensure all inputs are validated when using child_process.exec()'
            },
            # Generic dangerous function patterns
            {
                'pattern': r'(system\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'system() function with user input can lead to command injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use safer alternatives like subprocess with shell=False'
            },
            {
                'pattern': r'(shell_exec\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'shell_exec() with user input can lead to command injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use safer alternatives and validate all inputs'
            },
            # Low severity patterns - potential command injection vectors
            {
                'pattern': r'([\'"][^\'\"]*(?:Command|command|ls|cat|grep|find|echo)[^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Command-like string construction with user input',
                'severity': 'low',
                'confidence': 0.5,
                'remediation': 'Validate user input and avoid command-like string construction'
            },
            {
                'pattern': r'(os\.environ\s*\[\s*[\'"][^\'\"]*[\'"]\s*\]\s*=\s*\w+)',
                'description': 'Environment variable assignment with user input',
                'severity': 'low',
                'confidence': 0.4,
                'remediation': 'Validate environment variable values before assignment'
            },
            {
                'pattern': r'([\'"][^\'\"]*[\'"]\s*\+\s*\w+\s*\+\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Multiple string concatenations with user input',
                'severity': 'low',
                'confidence': 0.4,
                'remediation': 'Use safer string formatting methods'
            },
            {
                'pattern': r'(\{\s*[\'"][^\'\"]*[\'"]\s*:\s*\w+\s*\})',
                'description': 'Dictionary construction with user input',
                'severity': 'low',
                'confidence': 0.3,
                'remediation': 'Validate dictionary values before assignment'
            },
            {
                'pattern': r'([\'"][^\'\"]*\/[^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Path construction with user input',
                'severity': 'low',
                'confidence': 0.4,
                'remediation': 'Use os.path.join() for safe path construction'
            },
            {
                'pattern': r'([\'"][^\'\"]*[\'"]\s*\+\s*\w+\s*\+\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+\s*\+\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Complex string concatenation with multiple user inputs',
                'severity': 'low',
                'confidence': 0.3,
                'remediation': 'Use safer string formatting methods and validate all inputs'
            }
        ]
        
        for i, line in enumerate(lines, 1):
            line_stripped = line.strip()
            if not line_stripped or line_stripped.startswith('#'):
                continue
                
            for pattern_info in patterns:
                pattern = pattern_info['pattern']
                matches = re.findall(pattern, line, re.IGNORECASE)
                
                if matches:
                    for match in matches:
                        if isinstance(match, tuple):
                            code_snippet = match[0] if match else line_stripped
                        else:
                            code_snippet = match
                        
                        vulnerability = CommandInjectionVulnerability(
                            line_number=i,
                            vulnerability_type='Command Injection',
                            description=pattern_info['description'],
                            severity=pattern_info['severity'],
                            code_snippet=code_snippet,
                            remediation=pattern_info['remediation'],
                            confidence=pattern_info['confidence'],
                            file_path=filename
                        )
                        self.vulnerabilities.append(vulnerability)
    
    def _scan_with_ast(self, code, filename):
        """Scan for command injection vulnerabilities using AST analysis"""
        try:
            tree = ast.parse(code)
            visitor = CommandInjectionASTVisitor(filename)
            visitor.visit(tree)
            self.vulnerabilities.extend(visitor.vulnerabilities)
        except SyntaxError:
            # If the code has syntax errors, skip AST analysis
            pass

class CommandInjectionASTVisitor(ast.NodeVisitor):
    def __init__(self, filename):
        self.vulnerabilities = []
        self.filename = filename
        
    def visit_Call(self, node):
        """Visit function calls and check for command injection patterns"""
        if isinstance(node.func, ast.Attribute):
            # Check for os.system, os.popen, subprocess.* calls
            if (isinstance(node.func.value, ast.Name) and 
                node.func.value.id in ['os', 'subprocess'] and
                node.func.attr in ['system', 'popen', 'call', 'run', 'Popen']):
                
                # Check if arguments contain user input
                for arg in node.args:
                    if self._contains_user_input(arg):
                        vulnerability = CommandInjectionVulnerability(
                            line_number=node.lineno,
                            vulnerability_type='Command Injection',
                            description=f'Potential command injection in {node.func.value.id}.{node.func.attr}()',
                            severity='high',
                            code_snippet=self._get_code_snippet(node),
                            remediation='Use subprocess with shell=False and validate all inputs',
                            confidence=0.8,
                            file_path=self.filename
                        )
                        self.vulnerabilities.append(vulnerability)
        
        elif isinstance(node.func, ast.Name):
            # Check for eval, exec, __import__ calls
            if node.func.id in ['eval', 'exec', '__import__']:
                for arg in node.args:
                    if self._contains_user_input(arg):
                        vulnerability = CommandInjectionVulnerability(
                            line_number=node.lineno,
                            vulnerability_type='Code Injection',
                            description=f'Potential code injection in {node.func.id}()',
                            severity='high',
                            code_snippet=self._get_code_snippet(node),
                            remediation=f'Avoid {node.func.id}() or use safe alternatives',
                            confidence=0.9,
                            file_path=self.filename
                        )
                        self.vulnerabilities.append(vulnerability)
        
        self.generic_visit(node)
    
    def visit_BinOp(self, node):
        """Visit binary operations for string concatenation with user input"""
        if isinstance(node.op, ast.Add):
            # Check for string concatenation that might be used in commands
            left_str = self._extract_string_value(node.left)
            if left_str and any(cmd in left_str.lower() for cmd in ['sudo', 'su', 'bash', 'sh', 'cmd', 'python', 'perl', 'ruby']):
                if self._contains_user_input(node.right):
                    vulnerability = CommandInjectionVulnerability(
                        line_number=node.lineno,
                        vulnerability_type='Command Injection',
                        description='String concatenation with user input for command construction',
                        severity='high',
                        code_snippet=self._get_code_snippet(node),
                        remediation='Use subprocess with shell=False and validate inputs',
                        confidence=0.7,
                        file_path=self.filename
                    )
                    self.vulnerabilities.append(vulnerability)
        
        self.generic_visit(node)
    
    def visit_JoinedStr(self, node):
        """Visit f-strings for potential command injection"""
        # Check f-strings that might contain command construction
        for value in node.values:
            if isinstance(value, ast.FormattedValue):
                if self._contains_user_input(value.value):
                    # Check if the f-string contains command-like patterns
                    for other_value in node.values:
                        if isinstance(other_value, ast.Constant) and isinstance(other_value.value, str):
                            if any(cmd in other_value.value.lower() for cmd in ['sudo', 'su', 'bash', 'sh', 'cmd', 'python']):
                                vulnerability = CommandInjectionVulnerability(
                                    line_number=node.lineno,
                                    vulnerability_type='Command Injection',
                                    description='F-string with user input used for command construction',
                                    severity='high',
                                    code_snippet=self._get_code_snippet(node),
                                    remediation='Use subprocess with shell=False and validate inputs',
                                    confidence=0.8,
                                    file_path=self.filename
                                )
                                self.vulnerabilities.append(vulnerability)
                                break
        
        self.generic_visit(node)
    
    def _contains_user_input(self, node):
        """Check if a node contains potential user input"""
        # Common patterns for user input
        user_input_patterns = [
            'request', 'input', 'raw_input', 'sys.argv', 'os.environ',
            'form', 'args', 'json', 'data', 'params', 'query',
            'get', 'post', 'put', 'delete', 'patch', 'values',
            'headers', 'cookies', 'session', 'files'
        ]
        
        if isinstance(node, ast.Name):
            return any(pattern in node.id.lower() for pattern in user_input_patterns)
        elif isinstance(node, ast.Attribute):
            return any(pattern in node.attr.lower() for pattern in user_input_patterns)
        elif isinstance(node, ast.Subscript):
            return self._contains_user_input(node.value)
        elif isinstance(node, ast.Call):
            if isinstance(node.func, ast.Attribute):
                return any(pattern in node.func.attr.lower() for pattern in user_input_patterns)
            elif isinstance(node.func, ast.Name):
                return any(pattern in node.func.id.lower() for pattern in user_input_patterns)
        
        return False
    
    def _extract_string_value(self, node):
        """Extract string value from AST node if it's a constant string"""
        if isinstance(node, ast.Constant) and isinstance(node.value, str):
            return node.value
        return None
    
    def _get_code_snippet(self, node):
        """Get a code snippet for the AST node"""
        return f"Line {node.lineno}: {ast.dump(node)[:100]}..."

def highlight_command_injection_vulnerabilities(code, vulnerabilities=None):
    """Highlight command injection vulnerabilities in code"""
    if vulnerabilities is None:
        detector = CommandInjectionDetector()
        vulnerabilities = detector.scan_file('temp_code.py')
    
    lines = code.split('\n')
    highlighted_lines = []
    
    # Create a mapping of line numbers to vulnerabilities
    vuln_by_line = {}
    for vuln in vulnerabilities:
        if vuln.line_number not in vuln_by_line:
            vuln_by_line[vuln.line_number] = []
        vuln_by_line[vuln.line_number].append(vuln)
    
    for i, line in enumerate(lines, 1):
        if i in vuln_by_line:
            # Highlight the line with vulnerability information
            vulns = vuln_by_line[i]
            vuln_info = []
            for vuln in vulns:
                vuln_info.append(f"[{vuln.severity.upper()}] {vuln.description}")
            
            highlighted_line = {
                'line_number': i,
                'code': line,
                'highlighted': True,
                'vulnerabilities': vuln_info,
                'severity': vulns[0].severity if vulns else 'medium'
            }
        else:
            highlighted_line = {
                'line_number': i,
                'code': line,
                'highlighted': False,
                'vulnerabilities': [],
                'severity': None
            }
        
        highlighted_lines.append(highlighted_line)
    
    return {
        'highlighted_code': highlighted_lines,
        'total_vulnerabilities': len(vulnerabilities),
        'vulnerability_summary': {
            'critical': len([v for v in vulnerabilities if v.severity == 'critical']),
            'high': len([v for v in vulnerabilities if v.severity == 'high']),
            'medium': len([v for v in vulnerabilities if v.severity == 'medium']),
            'low': len([v for v in vulnerabilities if v.severity == 'low'])
        }
    }

def highlight_command_injection_vulnerabilities_word(code):
    """Generate highlighted code for Word document"""
    detector = CommandInjectionDetector()
    vulnerabilities = detector.scan_file('temp_code.py')
    
    lines = code.split('\n')
    highlighted_lines = []
    
    # Create a mapping of line numbers to vulnerabilities
    vuln_by_line = {}
    for vuln in vulnerabilities:
        if vuln.line_number not in vuln_by_line:
            vuln_by_line[vuln.line_number] = []
        vuln_by_line[vuln.line_number].append(vuln)
    
    for i, line in enumerate(lines, 1):
        if i in vuln_by_line:
            vulns = vuln_by_line[i]
            severity = vulns[0].severity if vulns else 'medium'
            highlighted_lines.append({
                'line_number': i,
                'code': line,
                'highlighted': True,
                'severity': severity,
                'vulnerabilities': [vuln.description for vuln in vulns]
            })
        else:
            highlighted_lines.append({
                'line_number': i,
                'code': line,
                'highlighted': False,
                'severity': None,
                'vulnerabilities': []
            })
    
    return highlighted_lines

def highlight_command_injection_vulnerabilities_html(code, vulnerabilities=None):
    """Highlight command injection vulnerabilities in code as HTML string (like XSS)"""
    if not vulnerabilities:
        return code
    lines = code.split('\n')
    highlighted_lines = []
    vuln_by_line = {}
    for vuln in vulnerabilities:
        line_num = vuln.line_number
        if line_num not in vuln_by_line:
            vuln_by_line[line_num] = []
        vuln_by_line[line_num].append(vuln)
    for line_num, line in enumerate(lines, 1):
        if line_num in vuln_by_line:
            line_vulns = vuln_by_line[line_num]
            severities = [v.severity for v in line_vulns]
            if 'critical' in severities:
                css_class = 'command-vuln-critical'
            elif 'high' in severities:
                css_class = 'command-vuln-high'
            elif 'medium' in severities:
                css_class = 'command-vuln-medium'
            else:
                css_class = 'command-vuln-low'
            patterns = []
            for vuln in line_vulns:
                snippet = vuln.code_snippet.strip()
                if snippet:
                    escaped_snippet = re.escape(snippet)
                    flexible_pattern = escaped_snippet.replace(r'\ ', r'\s*').replace(r'\"', r'["]').replace(r"\'", r'[\'"]')
                    patterns.append(flexible_pattern)
            highlighted_line = line
            for pattern in patterns:
                try:
                    highlighted_line = re.sub(f'({pattern})',
                        lambda m: f'<span class="{css_class}">{m.group(0)}</span>',
                        highlighted_line, flags=re.IGNORECASE)
                except re.error:
                    highlighted_line = f'<span class="{css_class}">{line}</span>'
                    break
            highlighted_lines.append(highlighted_line)
        else:
            highlighted_lines.append(line)
    return '\n'.join(highlighted_lines)

def scan_code_content_for_command_injection(code_content: str, source_name: str) -> dict:
    """Scan code content for command injection vulnerabilities"""
    # Create a temporary file to scan
    with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False) as temp_file:
        temp_file.write(code_content)
        temp_file_path = temp_file.name
    try:
        # Scan the temporary file
        detector = CommandInjectionDetector()
        vulnerabilities = detector.scan_file(temp_file_path)
        vuln_dicts = [vuln.to_dict() for vuln in vulnerabilities]
        # Get highlighted code as HTML string (like XSS)
        highlighted_code = highlight_command_injection_vulnerabilities_html(code_content, vulnerabilities)
        # Categorize vulnerabilities by severity
        severity_counts = {
            'high': 0,
            'medium': 0,
            'low': 0
        }
        for vuln in vulnerabilities:
            severity = vuln.severity.lower()
            if severity in severity_counts:
                severity_counts[severity] += 1
        total_vulnerabilities = len(vulnerabilities)
        risk_level = 'low'
        if severity_counts['high'] > 0:
            risk_level = 'high'
        elif severity_counts['medium'] > 0:
            risk_level = 'medium'
        return {
            'success': True,
            'source': source_name,
            'total_vulnerabilities': total_vulnerabilities,
            'total_issues': total_vulnerabilities,  # Frontend expects this
            'risk_level': risk_level,
            'severity_breakdown': severity_counts,
            # Frontend expects these fields
            'high_severity': severity_counts['high'],
            'medium_severity': severity_counts['medium'],
            'low_severity': severity_counts['low'],
            'high_count': severity_counts['high'],
            'medium_count': severity_counts['medium'],
            'low_count': severity_counts['low'],
            # Summary object for backward compatibility
            'summary': {
                'total_vulnerabilities': total_vulnerabilities,
                'high_severity': severity_counts['high'],
                'medium_severity': severity_counts['medium'],
                'low_severity': severity_counts['low'],
                'high': severity_counts['high'],
                'medium': severity_counts['medium'],
                'low': severity_counts['low']
            },
            'vulnerabilities': vuln_dicts,
            'highlighted_code': highlighted_code,
            'original_code': code_content,
            'scan_timestamp': datetime.now().isoformat(),
            'scanner_info': {
                'name': 'Command Injection Scanner',
                'version': '1.0.0',
                'description': 'Detects command injection vulnerabilities in Python code'
            }
        }
    finally:
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)

def is_github_py_url(url):
    """Check if URL is a GitHub Python file URL"""
    return 'github.com' in url and url.endswith('.py')

def github_raw_url(url):
    """Convert GitHub URL to raw URL"""
    return url.replace('github.com', 'raw.githubusercontent.com').replace('/blob/', '/')

def api_scan_command_injection(current_user):
    """API endpoint for command injection scanning"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        code_content = data.get('code')
        url = data.get('url')
        
        if code_content:
            # Scan provided code content
            result = scan_code_content_for_command_injection(code_content, 'Direct Input')
            
        elif url:
            # Scan URL content
            if not is_github_py_url(url):
                return jsonify({'error': 'Only GitHub Python files are supported'}), 400
            
            # Convert to raw URL
            raw_url = github_raw_url(url)
            
            # Fetch the code
            try:
                response = requests.get(raw_url, timeout=10)
                response.raise_for_status()
                code_content = response.text
                result = scan_code_content_for_command_injection(code_content, url)
            except requests.exceptions.RequestException as e:
                return jsonify({'error': f'Failed to fetch URL: {str(e)}'}), 400
        else:
            return jsonify({'error': 'No code content or URL provided'}), 400
        
        # Store result for report generation
        result_file = f'/tmp/results/command_injection_scan_{current_user}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json'
        with open(result_file, 'w') as f:
            json.dump(result, f, indent=2)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Scan failed: {str(e)}'}), 500

def api_generate_command_injection_report(current_user):
    """Generate Word report for command injection vulnerabilities"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        scan_result = data.get('scan_result')
        if not scan_result:
            return jsonify({'error': 'Scan result is required'}), 400
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Command Injection Vulnerability Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add executive summary
        doc.add_heading('Executive Summary', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Scan Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        summary_para.add_run(f"Source: {scan_result.get('source', 'Unknown')}\n")
        summary_para.add_run(f"Total Vulnerabilities: {scan_result.get('total_vulnerabilities', 0)}\n")
        summary_para.add_run(f"Risk Level: {scan_result.get('risk_level', 'Unknown').upper()}\n")
        
        # Add severity breakdown
        doc.add_heading('Severity Breakdown', level=1)
        severity_breakdown = scan_result.get('severity_breakdown', {})
        
        # Create table for severity breakdown
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Severity'
        hdr_cells[1].text = 'Count'
        
        for severity, count in severity_breakdown.items():
            row_cells = table.add_row().cells
            row_cells[0].text = severity.capitalize()
            row_cells[1].text = str(count)
        
        # Add vulnerabilities details
        if scan_result.get('vulnerabilities'):
            doc.add_heading('Vulnerability Details', level=1)
            
            for i, vuln in enumerate(scan_result['vulnerabilities'], 1):
                doc.add_heading(f'Vulnerability #{i}', level=2)
                
                # Create table for vulnerability details
                vuln_table = doc.add_table(rows=8, cols=2)
                vuln_table.style = 'Table Grid'
                
                details = [
                    ('Line Number', str(vuln.get('line_number', 'N/A'))),
                    ('Type', vuln.get('vulnerability_type', 'N/A')),
                    ('Severity', vuln.get('severity', 'N/A').upper()),
                    ('Confidence', f"{vuln.get('confidence', 0)*100:.1f}%"),
                    ('Description', vuln.get('description', 'N/A')),
                    ('Code Snippet', vuln.get('code_snippet', 'N/A')),
                    ('Remediation', vuln.get('remediation', 'N/A')),
                    ('References', f"CWE: {', '.join(vuln.get('cwe_references', []))}")
                ]
                
                for j, (key, value) in enumerate(details):
                    row = vuln_table.rows[j]
                    row.cells[0].text = key
                    row.cells[1].text = str(value)
                    
                    # Color code by severity
                    if key == 'Severity':
                        severity_color = {
                            'CRITICAL': RGBColor(255, 0, 0),
                            'HIGH': RGBColor(255, 165, 0),
                            'MEDIUM': RGBColor(255, 255, 0),
                            'LOW': RGBColor(0, 255, 0)
                        }
                        if value in severity_color:
                            row.cells[1].paragraphs[0].runs[0].font.color.rgb = severity_color[value]
        
        # Add highlighted code
        if scan_result.get('highlighted_code'):
            doc.add_heading('Code Analysis', level=1)
            doc.add_paragraph('The following code has been analyzed for command injection vulnerabilities:')
            
            # Add code with highlighting
            code_para = doc.add_paragraph()
            code_para.style = 'Code'
            
            for line_info in scan_result['highlighted_code']:
                line_text = f"{line_info['line_number']:3d}: {line_info['code']}\n"
                run = code_para.add_run(line_text)
                
                if line_info.get('highlighted'):
                    # Highlight vulnerable lines
                    severity = line_info.get('severity', 'medium')
                    if severity == 'critical':
                        run.font.highlight_color = WD_COLOR_INDEX.RED
                    elif severity == 'high':
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    elif severity == 'medium':
                        run.font.highlight_color = WD_COLOR_INDEX.CYAN
                    else:
                        run.font.highlight_color = WD_COLOR_INDEX.GREEN
        
        # Add recommendations
        doc.add_heading('Recommendations', level=1)
        recommendations = [
            "Use subprocess.run() with shell=False instead of os.system() or os.popen()",
            "Validate and sanitize all user inputs before using them in commands",
            "Use parameterized commands with proper argument lists",
            "Implement input validation and whitelisting for allowed commands",
            "Use secure coding practices and regular security reviews",
            "Consider using higher-level APIs that don't expose shell execution",
            "Implement proper logging and monitoring for command execution"
        ]
        
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Bullet')
        
        # Save document
        doc_filename = f'/tmp/results/command_injection_report_{current_user}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(doc_filename)
        
        return send_file(doc_filename, as_attachment=True, download_name='command_injection_report.docx')
        
    except Exception as e:
        return jsonify({'error': f'Report generation failed: {str(e)}'}), 500

def api_command_injection_sonarqube_export(current_user):
    """Export command injection vulnerabilities in SonarQube format"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        scan_result = data.get('scan_result')
        if not scan_result:
            return jsonify({'error': 'Scan result is required'}), 400
        
        # Create SonarQube-compatible issue export
        issues = []
        
        for vuln in scan_result.get('vulnerabilities', []):
            issue = {
                "engineId": "command-injection-scanner",
                "ruleId": vuln.get('rule_key', 'python:S2076'),
                "severity": vuln.get('severity', 'MEDIUM').upper(),
                "type": "VULNERABILITY",
                "primaryLocation": {
                    "message": vuln.get('description', 'Command injection vulnerability detected'),
                    "filePath": vuln.get('file_path', 'unknown'),
                    "textRange": {
                        "startLine": vuln.get('line_number', 1),
                        "endLine": vuln.get('line_number', 1),
                        "startColumn": 1,
                        "endColumn": len(vuln.get('code_snippet', ''))
                    }
                },
                "effortMinutes": 30,
                "tags": ["command-injection", "security", "injection"],
                "cwe": vuln.get('cwe_references', []),
                "owaspTop10": vuln.get('owasp_references', []),
                "cleanCodeAttribute": "TRUSTWORTHY",
                "impacts": [
                    {
                        "softwareQuality": "SECURITY",
                        "severity": "HIGH" if vuln.get('severity', 'medium').upper() in ['CRITICAL', 'HIGH'] else "MEDIUM"
                    }
                ]
            }
            issues.append(issue)
        
        # Create the export structure
        export_data = {
            "issues": issues,
            "metadata": {
                "toolName": "Command Injection Scanner",
                "toolVersion": "1.0.0",
                "scanDate": datetime.now().isoformat(),
                "source": scan_result.get('source', 'Unknown'),
                "totalIssues": len(issues),
                "riskLevel": scan_result.get('risk_level', 'medium')
            }
        }
        
        # Save export file
        export_filename = f'/tmp/results/command_injection_sonarqube_export_{current_user}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json'
        with open(export_filename, 'w') as f:
            json.dump(export_data, f, indent=2)
        
        return send_file(export_filename, as_attachment=True, download_name='command_injection_sonarqube_export.json')
        
    except Exception as e:
        return jsonify({'error': f'Export failed: {str(e)}'}), 500 