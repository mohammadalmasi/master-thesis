from flask import request, jsonify, send_file
import requests
import re
import os
import ast
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from sonarqube_security_standards import SecurityStandards, SQCategory, VulnerabilityProbability
from datetime import datetime
import json
import tempfile

# Based on SonarQube's SecurityStandards.java
XSS_VULNERABILITY = ("xss", VulnerabilityProbability.HIGH)
# Maps to CWE-79, CWE-80, CWE-81, CWE-82, CWE-83, CWE-84, CWE-85, CWE-86, CWE-87
# Maps to OWASP A03:2021-Injection

class XSSVulnerability:
    def __init__(self, line_number, vulnerability_type, description, severity, code_snippet, remediation, confidence, file_path=None):
        self.line_number = line_number
        self.vulnerability_type = vulnerability_type
        self.description = description
        self.severity = severity
        self.code_snippet = code_snippet
        self.remediation = remediation
        self.confidence = confidence
        self.file_path = file_path or 'unknown'
        
    def to_dict(self):
        return {
            'line_number': self.line_number,
            'vulnerability_type': self.vulnerability_type,
            'description': self.description,
            'severity': self.severity,
            'code_snippet': self.code_snippet,
            'remediation': self.remediation,
            'confidence': self.confidence,
            'file_path': self.file_path,
            'cwe_references': ["79", "80", "81", "82", "83", "84", "85", "86", "87"],
            'owasp_references': ["A03:2021-Injection"],
            'rule_key': 'python:S5131'  # XSS rule key similar to SonarQube
        }

class XSSDetector:
    def __init__(self):
        self.vulnerabilities = []
        
    def scan_file(self, filename):
        """Scan a Python file for XSS vulnerabilities"""
        self.vulnerabilities = []
        
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                code = f.read()
        except UnicodeDecodeError:
            with open(filename, 'r', encoding='latin-1') as f:
                code = f.read()
        
        # Scan for XSS patterns
        self._scan_for_patterns(code, filename)
        self._scan_with_ast(code, filename)
        
        # Remove duplicates
        self.vulnerabilities = self._deduplicate_vulnerabilities(self.vulnerabilities)
        
        return self.vulnerabilities
    
    def _deduplicate_vulnerabilities(self, vulnerabilities):
        """Remove duplicate vulnerabilities based on line number and similarity"""
        if not vulnerabilities:
            return vulnerabilities
        
        # Group vulnerabilities by line number
        by_line = {}
        for vuln in vulnerabilities:
            line_key = vuln.line_number
            if line_key not in by_line:
                by_line[line_key] = []
            by_line[line_key].append(vuln)
        
        # For each line, keep only the best vulnerability
        deduplicated = []
        for line_num, line_vulns in by_line.items():
            if len(line_vulns) == 1:
                deduplicated.extend(line_vulns)
            else:
                # Multiple vulnerabilities on the same line - keep the best one
                # Sort by confidence (highest first), then by description length (longest first)
                best_vuln = max(line_vulns, key=lambda v: (v.confidence, len(v.description)))
                deduplicated.append(best_vuln)
        
        return deduplicated
    
    def _scan_for_patterns(self, code, filename):
        """Scan for XSS vulnerability patterns using regex"""
        lines = code.split('\n')
        
        # XSS vulnerability patterns
        patterns = [
            # Direct HTML output without escaping
            {
                'pattern': r'(render_template_string\s*\(\s*[\'"][^\'\"]*\{\{[^}]*\}\}[^\'\"]*[\'"][^)]*\))',
                'description': 'Direct template rendering with user input can lead to XSS',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use safe template rendering with proper escaping or template files'
            },
            # innerHTML patterns - both concatenation and direct assignment
            {
                'pattern': r'(\.innerHTML\s*=\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Direct innerHTML manipulation with concatenated user input',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use textContent or properly escape HTML content'
            },
            {
                'pattern': r'(\.innerHTML\s*=\s*(?![\'"]).+)',
                'description': 'Direct innerHTML assignment with potentially unsafe content',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use textContent or properly escape HTML content before assignment'
            },
            {
                'pattern': r'(document\.getElementById\([^)]+\)\.innerHTML\s*=\s*[^;]+)',
                'description': 'Direct DOM manipulation via innerHTML - potential XSS vector',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use textContent or sanitize content before innerHTML assignment'
            },
            # URL parameter usage patterns
            {
                'pattern': r'(URLSearchParams|location\.search|params\.get|getParameter)',
                'description': 'URL parameter access detected - ensure proper sanitization if used in DOM',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Sanitize URL parameters before using in DOM manipulation or HTML output'
            },
            {
                'pattern': r'(document\.write\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'document.write() with user input can lead to XSS',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Avoid document.write() or properly escape user input'
            },
            {
                'pattern': r'(document\.write\s*\([^)]*\w+[^)]*\))',
                'description': 'document.write() with potentially unsafe content',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Avoid document.write() or properly escape user input'
            },
            {
                'pattern': r'(eval\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'eval() with user input can lead to code injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Avoid eval() or use safe alternatives like JSON.parse()'
            },
            {
                'pattern': r'(eval\s*\([^)]*\w+[^)]*\))',
                'description': 'eval() with potentially unsafe content',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Avoid eval() completely or use safe alternatives like JSON.parse()'
            },
            # DOM manipulation patterns
            {
                'pattern': r'(\.outerHTML\s*=\s*[^;]+)',
                'description': 'Direct outerHTML manipulation can lead to XSS',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use safer DOM manipulation methods or sanitize content'
            },
            {
                'pattern': r'(\.insertAdjacentHTML\s*\([^)]*\))',
                'description': 'insertAdjacentHTML with unsanitized content can lead to XSS',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Sanitize HTML content or use textContent/createElement instead'
            },
            # jQuery patterns
            {
                'pattern': r'(\$\([^)]*\)\.html\s*\([^)]*\w+[^)]*\))',
                'description': 'jQuery .html() with potentially unsafe content',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use .text() or sanitize content before using .html()'
            },
            {
                'pattern': r'(\$\([^)]*\)\.append\s*\([^)]*<[^>]*>[^)]*\))',
                'description': 'jQuery .append() with HTML content - potential XSS',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Sanitize HTML content or use text-only methods'
            },
            # Flask/Jinja2 specific patterns
            {
                'pattern': r'(\{\{\s*\w+\s*\|\s*safe\s*\}\})',
                'description': 'Using |safe filter without proper validation can lead to XSS',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Ensure input is properly validated and sanitized before using |safe filter'
            },
            {
                'pattern': r'(Markup\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Markup() with concatenated user input can lead to XSS',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use Markup() only with trusted content or properly escape user input'
            },
            {
                'pattern': r'(Markup\s*\([^)]*\w+[^)]*\))',
                'description': 'Markup() with potentially unsafe content',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Ensure content is trusted and sanitized before using Markup()'
            },
            # Direct output patterns
            {
                'pattern': r'(print\s*\(\s*[\'"]<[^>]*>[\'\"]\s*\+\s*\w+)',
                'description': 'Printing HTML with user input can lead to XSS in web contexts',
                'severity': 'medium',
                'confidence': 0.6,
                'remediation': 'Escape HTML content or use template engines'
            },
            # F-string patterns with HTML content
            {
                'pattern': r'(f[\'"]<[^>]*>[^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
                'description': 'F-string with HTML content and user input can lead to XSS',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use template engines or escape HTML characters before embedding in f-strings'
            },
            # Return statements with HTML f-strings
            {
                'pattern': r'(return\s+f[\'"]<[^>]*>[^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
                'description': 'Returning HTML f-string with user input can lead to XSS',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use template engines or escape HTML characters before returning'
            },
            {
                'pattern': r'(response\.write\s*\(\s*[\'"]<[^>]*>[\'\"]\s*\+\s*\w+)',
                'description': 'Writing HTML response with user input can lead to XSS',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Escape HTML content before writing to response'
            },
            # Request parameter usage in HTML context
            {
                'pattern': r'(request\.args\.get\([^)]*\)[^;]*<[^>]*>)',
                'description': 'Using request parameters directly in HTML context',
                'severity': 'high',
                'confidence': 0.7,
                'remediation': 'Escape HTML characters in user input'
            },
            {
                'pattern': r'(request\.form\.get\([^)]*\)[^;]*<[^>]*>)',
                'description': 'Using form data directly in HTML context',
                'severity': 'high',
                'confidence': 0.7,
                'remediation': 'Escape HTML characters in user input'
            },
            # String formatting with HTML
            {
                'pattern': r'([\'"]<[^>]*>[\'\"]\s*%\s*\w+)',
                'description': 'String formatting with user input in HTML context',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use template engines or escape HTML characters'
            },
            {
                'pattern': r'([\'"]<[^>]*>.*\{\}.*[\'\"]\s*\.format\s*\([^)]*\))',
                'description': 'Format string with user input in HTML context',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use template engines or escape HTML characters'
            },
            # JavaScript context
            {
                'pattern': r'(<script[^>]*>[^<]*\+\s*\w+[^<]*</script>)',
                'description': 'Direct user input in JavaScript context can lead to XSS',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Properly encode data for JavaScript context or use JSON'
            },
            {
                'pattern': r'(<script[^>]*>.*innerHTML.*</script>)',
                'description': 'innerHTML usage in script tags - potential XSS vector',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use textContent or properly sanitize content'
            },
            # Additional dangerous functions
            {
                'pattern': r'(setTimeout\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'setTimeout with string concatenation can lead to code injection',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use function references instead of string evaluation'
            },
            {
                'pattern': r'(setInterval\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'setInterval with string concatenation can lead to code injection',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use function references instead of string evaluation'
            },
            # LOW SEVERITY PATTERNS
            # Simple string concatenation without HTML tags
            {
                'pattern': r'(\w+\s*\+\s*[\'"][^<>]*[\'"])',
                'description': 'Simple string concatenation - potential low-risk XSS',
                'severity': 'low',
                'confidence': 0.3,
                'remediation': 'Validate input if used in HTML context'
            },
            # URL parameter access (flagged for awareness)
            {
                'pattern': r'(URLSearchParams|location\.search|params\.get|getParameter)',
                'description': 'URL parameter access detected - ensure proper sanitization if used in DOM',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Sanitize URL parameters before using in DOM manipulation or HTML output'
            },
            # Template string building
            {
                'pattern': r'(\w+\s*=\s*[\'"][^<>]*[\'\"]\s*\+\s*\w+)',
                'description': 'String template building - low risk if not output to HTML',
                'severity': 'low',
                'confidence': 0.4,
                'remediation': 'Ensure proper escaping if used in HTML context'
            }
        ]
        
        for line_num, line in enumerate(lines, 1):
            for pattern_info in patterns:
                matches = re.finditer(pattern_info['pattern'], line, re.IGNORECASE)
                for match in matches:
                    vulnerability = XSSVulnerability(
                        line_number=line_num,
                        vulnerability_type='xss',
                        description=pattern_info['description'],
                        severity=pattern_info['severity'],
                        code_snippet=line.strip(),
                        remediation=pattern_info['remediation'],
                        confidence=pattern_info['confidence'],
                        file_path=filename
                    )
                    self.vulnerabilities.append(vulnerability)
    
    def _scan_with_ast(self, code, filename):
        """Scan for XSS vulnerabilities using AST analysis"""
        try:
            tree = ast.parse(code)
            visitor = XSSASTVisitor(filename)
            visitor.visit(tree)
            self.vulnerabilities.extend(visitor.vulnerabilities)
        except SyntaxError:
            # If code has syntax errors, skip AST analysis
            pass

class XSSASTVisitor(ast.NodeVisitor):
    def __init__(self, filename):
        self.filename = filename
        self.vulnerabilities = []
    
    def visit_Call(self, node):
        """Visit function calls to detect XSS patterns"""
        if isinstance(node.func, ast.Attribute):
            # Check for dangerous methods
            if (hasattr(node.func, 'attr') and 
                node.func.attr in ['render_template_string', 'write', 'send']):
                
                # Check if arguments contain user input patterns
                for arg in node.args:
                    if self._contains_user_input(arg):
                        vulnerability = XSSVulnerability(
                            line_number=getattr(node, 'lineno', 0),
                            vulnerability_type='xss',
                            description=f'Potentially unsafe {node.func.attr}() call with user input',
                            severity='high',
                            code_snippet=self._get_code_snippet(node),
                            remediation='Ensure proper input validation and output encoding',
                            confidence=0.7,
                            file_path=self.filename
                        )
                        self.vulnerabilities.append(vulnerability)
        
        self.generic_visit(node)
    
    def visit_BinOp(self, node):
        """Visit binary operations for string concatenation in HTML context"""
        if isinstance(node.op, ast.Add):
            # Check for HTML-like patterns in string concatenation
            left_str = self._extract_string_value(node.left)
            right_str = self._extract_string_value(node.right)
            
            if ((left_str and '<' in left_str and '>' in left_str) or
                (right_str and '<' in right_str and '>' in right_str)):
                
                if self._contains_user_input(node.right) or self._contains_user_input(node.left):
                    vulnerability = XSSVulnerability(
                        line_number=getattr(node, 'lineno', 0),
                        vulnerability_type='xss',
                        description='String concatenation with HTML content and potential user input',
                        severity='medium',
                        code_snippet=self._get_code_snippet(node),
                        remediation='Use template engines or escape HTML characters',
                        confidence=0.6,
                        file_path=self.filename
                    )
                    self.vulnerabilities.append(vulnerability)
        
        self.generic_visit(node)
    
    def visit_JoinedStr(self, node):
        """Visit f-strings (JoinedStr nodes) for XSS patterns"""
        # Check if f-string contains HTML-like content
        html_pattern = False
        has_user_input = False
        
        for value in node.values:
            if isinstance(value, ast.Str):
                if '<' in value.s and '>' in value.s:
                    html_pattern = True
            elif isinstance(value, ast.Constant) and isinstance(value.value, str):
                if '<' in value.value and '>' in value.value:
                    html_pattern = True
            elif isinstance(value, ast.FormattedValue):
                # Check if the formatted value contains user input
                if self._contains_user_input(value.value):
                    has_user_input = True
        
        if html_pattern and has_user_input:
            vulnerability = XSSVulnerability(
                line_number=getattr(node, 'lineno', 0),
                vulnerability_type='xss',
                description='F-string with HTML content and user input can lead to XSS',
                severity='high',
                code_snippet=self._get_code_snippet(node),
                remediation='Use template engines or escape HTML characters before embedding in f-strings',
                confidence=0.9,
                file_path=self.filename
            )
            self.vulnerabilities.append(vulnerability)
        
        self.generic_visit(node)
    
    def _contains_user_input(self, node):
        """Check if node contains patterns that might be user input"""
        if isinstance(node, ast.Attribute):
            if (hasattr(node, 'attr') and 
                node.attr in ['args', 'form', 'cookies', 'headers', 'json']):
                return True
            # Check for request.args patterns
            if (isinstance(node.value, ast.Attribute) and
                hasattr(node.value, 'attr') and
                node.value.attr == 'args' and
                isinstance(node.value.value, ast.Name) and
                node.value.value.id == 'request'):
                return True
        elif isinstance(node, ast.Call):
            if (isinstance(node.func, ast.Attribute) and
                hasattr(node.func, 'attr') and
                node.func.attr in ['get', 'getlist']):
                return True
            # Check for request.args.get() patterns
            if (isinstance(node.func, ast.Attribute) and
                hasattr(node.func, 'attr') and
                node.func.attr == 'get' and
                isinstance(node.func.value, ast.Attribute) and
                hasattr(node.func.value, 'attr') and
                node.func.value.attr in ['args', 'form', 'cookies', 'headers', 'json']):
                return True
        elif isinstance(node, ast.Name):
            # Check for variables that might contain user input
            if hasattr(node, 'id') and node.id in ['name', 'username', 'input', 'data', 'param', 'value']:
                return True
        return False
    
    def _extract_string_value(self, node):
        """Extract string value from AST node if it's a string literal"""
        if isinstance(node, ast.Str):
            return node.s
        elif isinstance(node, ast.Constant) and isinstance(node.value, str):
            return node.value
        return None
    
    def _get_code_snippet(self, node):
        """Get a string representation of the AST node"""
        try:
            return ast.unparse(node)
        except AttributeError:
            # Fallback for Python < 3.9
            return f'Line {getattr(node, "lineno", "unknown")}'

def highlight_xss_vulnerabilities(code, vulnerabilities=None):
    """Highlight XSS vulnerability patterns in code with severity-based colors"""
    if not vulnerabilities:
        return code
    
    # Split code into lines for line-by-line highlighting
    lines = code.split('\n')
    highlighted_lines = []
    
    # Group vulnerabilities by line number
    vuln_by_line = {}
    for vuln in vulnerabilities:
        line_num = vuln.line_number
        if line_num not in vuln_by_line:
            vuln_by_line[line_num] = []
        vuln_by_line[line_num].append(vuln)
    
    # Process each line
    for line_num, line in enumerate(lines, 1):
        if line_num in vuln_by_line:
            # Find the highest severity for this line
            line_vulns = vuln_by_line[line_num]
            severities = [v.severity for v in line_vulns]
            
            # Determine CSS class based on highest severity
            if 'high' in severities:
                css_class = 'xss-vuln-high'
            elif 'medium' in severities:
                css_class = 'xss-vuln-medium'
            else:
                css_class = 'xss-vuln-low'
            
            # Create detailed patterns for this line based on actual vulnerabilities
            patterns = []
            for vuln in line_vulns:
                # Create pattern from the vulnerable code snippet
                snippet = vuln.code_snippet.strip()
                if snippet:
                    # Escape special regex characters and create a flexible pattern
                    escaped_snippet = re.escape(snippet)
                    # Make it more flexible to catch variations
                    flexible_pattern = escaped_snippet.replace(r'\ ', r'\s*').replace(r'\"', r'["\']').replace(r'\'', r'["\']')
                    patterns.append(flexible_pattern)
            
            # Apply highlighting to this line
            highlighted_line = line
            for pattern in patterns:
                try:
                    highlighted_line = re.sub(f'({pattern})', 
                                            lambda m: f'<span class="{css_class}">{m.group(0)}</span>', 
                                            highlighted_line, flags=re.IGNORECASE)
                except re.error:
                    # If regex fails, highlight the whole line
                    highlighted_line = f'<span class="{css_class}">{line}</span>'
                    break
            
            highlighted_lines.append(highlighted_line)
        else:
            highlighted_lines.append(line)
    
    return '\n'.join(highlighted_lines)

def highlight_xss_vulnerabilities_word(code):
    """Highlight XSS vulnerability patterns for Word documents"""
    patterns = [
        r'(render_template_string\s*\([^)]*\))',
        r'(\.innerHTML\s*=\s*[^;]*)',
        r'(document\.getElementById\([^)]+\)\.innerHTML\s*=\s*[^;]+)',
        r'(\.outerHTML\s*=\s*[^;]+)',
        r'(\.insertAdjacentHTML\s*\([^)]*\))',
        r'(document\.write\s*\([^)]*\))',
        r'(eval\s*\([^)]*\))',
        r'(\$\([^)]*\)\.html\s*\([^)]*\))',
        r'(\$\([^)]*\)\.append\s*\([^)]*<[^>]*>[^)]*\))',
        r'(\{\{\s*\w+\s*\|\s*safe\s*\}\})',
        r'(Markup\s*\([^)]*\))',
        r'(URLSearchParams|location\.search|params\.get|getParameter)',
        r'(request\.args\.get\([^)]*\))',
        r'(request\.form\.get\([^)]*\))',
        r'(<script[^>]*>[^<]*</script>)',
        r'([\'"]<[^>]*>.*\{\}.*[\'\"]\s*\.format\s*\([^)]*\))',
        r'(f[\'"]<[^>]*>[^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'(return\s+f[\'"]<[^>]*>[^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'(setTimeout\s*\([^)]*\))',
        r'(setInterval\s*\([^)]*\))'
    ]
    
    highlighted = code
    for pattern in patterns:
        highlighted = re.sub(pattern, lambda m: f'[XSS-VULNERABLE:{m.group(0)}]', highlighted, flags=re.IGNORECASE)
    
    return highlighted

def scan_code_content_for_xss(code_content: str, source_name: str) -> dict:
    """Scan code content for XSS vulnerabilities"""
    try:
        # Create temporary file for scanning
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False)
        temp_file.write(code_content)
        temp_file.close()
        
        # Create XSS detector and scan
        detector = XSSDetector()
        vulnerabilities = detector.scan_file(temp_file.name)
        
        # Clean up temporary file
        os.unlink(temp_file.name)
        
        # Generate highlighted code
        highlighted_code = None
        original_code = code_content
        file_name = source_name
        
        if vulnerabilities:
            highlighted_code = highlight_xss_vulnerabilities(code_content, vulnerabilities)
            original_code = code_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            if '/' in source_name:
                file_name = source_name.split('/')[-1]
            elif source_name.startswith('http'):
                file_name = source_name.split('/')[-1] if '/' in source_name else 'scanned_code.py'
        
        # Calculate summary
        summary = {
            'total_vulnerabilities': len(vulnerabilities),
            'high_severity': sum(1 for v in vulnerabilities if v.severity == 'high'),
            'medium_severity': sum(1 for v in vulnerabilities if v.severity == 'medium'),
            'low_severity': sum(1 for v in vulnerabilities if v.severity == 'low'),
            'high': sum(1 for v in vulnerabilities if v.severity == 'high'),
            'medium': sum(1 for v in vulnerabilities if v.severity == 'medium'),
            'low': sum(1 for v in vulnerabilities if v.severity == 'low')
        }
        
        # Format results
        results = {
            'source': source_name,
            'scan_type': 'xss',
            'summary': summary,
            'vulnerabilities': [vuln.to_dict() for vuln in vulnerabilities],
            'total_vulnerabilities': len(vulnerabilities),
            'scan_timestamp': datetime.now().isoformat(),
            'total_issues': len(vulnerabilities),
            'high_severity': summary['high_severity'],
            'medium_severity': summary['medium_severity'],
            'low_severity': summary['low_severity'],
            'high_count': summary['high'],
            'medium_count': summary['medium'],
            'low_count': summary['low'],
            'highlighted_code': highlighted_code,
            'original_code': original_code,
            'file_name': file_name
        }
        
        return results
        
    except Exception as e:
        return {
            'error': f'Error during XSS scan: {str(e)}',
            'source': source_name,
            'scan_type': 'xss',
            'vulnerabilities': [],
            'total_vulnerabilities': 0,
            'total_issues': 0,
            'high_severity': 0,
            'medium_severity': 0,
            'low_severity': 0,
            'high_count': 0,
            'medium_count': 0,
            'low_count': 0,
            'highlighted_code': None,
            'original_code': '',
            'file_name': source_name
        }

def is_github_py_url(url):
    """Check if URL is a GitHub Python file URL"""
    return 'github.com' in url and url.endswith('.py')

def github_raw_url(url):
    """Convert GitHub blob URL to raw URL"""
    return url.replace('github.com', 'raw.githubusercontent.com').replace('/blob/', '/')

def api_scan_xss(current_user):
    """API endpoint for XSS scanning"""
    try:
        data = request.get_json()
        code_content = data.get('code')
        url = data.get('url')
        
        if code_content:
            # Scan provided code content
            results = scan_code_content_for_xss(code_content, 'Direct input')
            
        elif url:
            # Scan URL content
            if is_github_py_url(url):
                raw_url = github_raw_url(url)
                try:
                    response = requests.get(raw_url, timeout=10)
                    if response.status_code == 200:
                        results = scan_code_content_for_xss(response.text, url)
                    else:
                        return jsonify({'error': f'Failed to fetch URL: {response.status_code}'}), 400
                except Exception as e:
                    return jsonify({'error': f'Error fetching URL: {str(e)}'}), 400
            else:
                return jsonify({'error': 'Invalid GitHub Python file URL'}), 400
        else:
            return jsonify({'error': 'Invalid scan parameters'}), 400
        
        return jsonify(results)
        
    except Exception as e:
        return jsonify({'error': f'Error during XSS scan: {str(e)}'}), 500

def api_generate_xss_report(current_user):
    """API endpoint for generating XSS reports"""
    try:
        data = request.get_json()
        vulnerabilities = data.get('vulnerabilities', [])
        source = data.get('source', 'Unknown')
        
        if not vulnerabilities:
            return jsonify({'error': 'No vulnerabilities provided'}), 400
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Cross-Site Scripting (XSS) Security Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_heading('Report Information', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        info_table.cell(0, 0).text = 'Source'
        info_table.cell(0, 1).text = source
        info_table.cell(1, 0).text = 'Report Generated'
        info_table.cell(1, 1).text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        info_table.cell(2, 0).text = 'Total Vulnerabilities'
        info_table.cell(2, 1).text = str(len(vulnerabilities))
        info_table.cell(3, 0).text = 'Risk Level'
        info_table.cell(3, 1).text = 'High' if len(vulnerabilities) > 0 else 'Low'
        
        # Add summary
        doc.add_heading('Executive Summary', level=1)
        summary_text = f"""
This report presents the results of a comprehensive Cross-Site Scripting (XSS) vulnerability analysis performed on the provided code. 
The analysis identified {len(vulnerabilities)} potential XSS vulnerabilities that could allow attackers to inject malicious scripts 
into web applications and execute them in users' browsers.

XSS vulnerabilities can lead to session hijacking, credential theft, malware distribution, and defacement of web applications.
These vulnerabilities should be addressed immediately to prevent potential attacks against users.
        """
        doc.add_paragraph(summary_text.strip())
        
        # Add vulnerability details
        doc.add_heading('Vulnerability Details', level=1)
        
        for i, vuln in enumerate(vulnerabilities, 1):
            doc.add_heading(f'XSS Vulnerability #{i}', level=2)
            
            # Vulnerability info table
            vuln_table = doc.add_table(rows=5, cols=2)
            vuln_table.style = 'Table Grid'
            
            vuln_table.cell(0, 0).text = 'Line Number'
            vuln_table.cell(0, 1).text = str(vuln.get('line_number', 'N/A'))
            vuln_table.cell(1, 0).text = 'Severity'
            vuln_table.cell(1, 1).text = vuln.get('severity', 'Medium').title()
            vuln_table.cell(2, 0).text = 'CWE References'
            vuln_table.cell(2, 1).text = ', '.join(vuln.get('cwe_references', []))
            vuln_table.cell(3, 0).text = 'OWASP References'
            vuln_table.cell(3, 1).text = ', '.join(vuln.get('owasp_references', []))
            vuln_table.cell(4, 0).text = 'Description'
            vuln_table.cell(4, 1).text = vuln.get('description', 'XSS vulnerability detected')
            
            # Vulnerable code
            doc.add_paragraph('Vulnerable Code:', style='Heading 3')
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(vuln.get('code_snippet', ''))
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(10)
            code_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
            
            # Remediation
            doc.add_paragraph('Remediation:', style='Heading 3')
            remediation_para = doc.add_paragraph(vuln.get('remediation', 'Apply proper input validation and output encoding'))
            
            doc.add_paragraph()  # Add spacing
        
        # Add recommendations
        doc.add_heading('XSS Prevention Recommendations', level=1)
        recommendations = """
1. Input Validation: Validate all user inputs on both client and server side.
2. Output Encoding: Encode data before inserting it into HTML, JavaScript, CSS, or URL contexts.
3. Content Security Policy (CSP): Implement strict CSP headers to prevent script execution.
4. Use Safe APIs: Avoid innerHTML, document.write(), and eval(). Use textContent, createElement(), and JSON.parse().
5. Template Engines: Use template engines with built-in XSS protection (auto-escaping).
6. HTTPOnly Cookies: Set HTTPOnly flag on sensitive cookies to prevent JavaScript access.
7. Security Headers: Implement X-XSS-Protection, X-Frame-Options, and X-Content-Type-Options headers.
8. Regular Testing: Conduct regular security testing including automated XSS scanning.
        """
        doc.add_paragraph(recommendations.strip())
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'xss_security_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': f'Error generating XSS report: {str(e)}'}), 500

def api_xss_sonarqube_export(current_user):
    """API endpoint for XSS SonarQube export"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        vulnerabilities = data.get('vulnerabilities', [])
        if not vulnerabilities:
            return jsonify({'error': 'No vulnerabilities provided'}), 400
        
        # Convert to SonarQube format
        sonar_issues = []
        for vuln in vulnerabilities:
            sonar_issue = {
                "engineId": "python-xss-scanner",
                "ruleId": vuln.get('rule_key', 'python:S5131'),
                "severity": vuln.get('severity', 'MAJOR').upper(),
                "type": "VULNERABILITY",
                "primaryLocation": {
                    "message": vuln.get('description', 'Cross-Site Scripting (XSS) vulnerability'),
                    "filePath": vuln.get('file_path', 'unknown'),
                    "textRange": {
                        "startLine": vuln.get('line_number', 1),
                        "endLine": vuln.get('line_number', 1)
                    }
                },
                "cwe": vuln.get('cwe_references', []),
                "owasp": vuln.get('owasp_references', []),
                "confidence": vuln.get('confidence', 0.5)
            }
            sonar_issues.append(sonar_issue)
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.json')
        json.dump({"issues": sonar_issues}, temp_file, indent=2)
        temp_file.close()
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'xss_sonarqube_issues_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json',
            mimetype='application/json'
        )
        
    except Exception as e:
        return jsonify({'error': f'Error exporting XSS SonarQube format: {str(e)}'}), 500 