from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import sqlite3
import json
from datetime import datetime, timedelta
# import jwt
# from functools import wraps
import tempfile
import zipfile
import uuid
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Import XSS scanner functions
from scanners.xss.xss_scanner import (
    api_scan_xss,
    api_generate_xss_report,
    api_xss_sonarqube_export
)

# Import SQL injection scanner functions  
from scanners.sql_injection.sql_injection_scanner import (
    api_scan_sql_injection,
    api_generate_sql_injection_report,
    api_sql_injection_sonarqube_export
)

# Import Command injection scanner functions
from scanners.command_injection.command_injection_scanner import (
    api_scan_command_injection,
    api_generate_command_injection_report,
    api_command_injection_sonarqube_export
)

# Import CSRF scanner functions
from scanners.csrf.csrf_scanner import (
    api_scan_csrf,
    api_generate_csrf_report,
    api_csrf_sonarqube_export
)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 2 * 1024 * 1024  # 2MB upload limit
# app.config['JWT_SECRET_KEY'] = os.environ.get('JWT_SECRET_KEY', 'sql-injection-scanner-secret-key-2024')

# Enable CORS for all origins (production and development)
CORS(app, origins=["*"], 
     methods=["GET", "POST", "OPTIONS"], 
     allow_headers=["Content-Type", "Authorization"])

# Authentication configuration
# ADMIN_USERNAME = "admin"
# ADMIN_PASSWORD = "a"

# def token_required(f):
#     @wraps(f)
#     def decorated(*args, **kwargs):
#         token = None
#         
#         # JWT is passed in the request header
#         if 'Authorization' in request.headers:
#             auth_header = request.headers['Authorization']
#             try:
#                 token = auth_header.split(" ")[1]  # Bearer <token>
#             except IndexError:
#                 return jsonify({'error': 'Token is missing!'}), 401
#         
#         if not token:
#             return jsonify({'error': 'Token is missing!'}), 401
#         
#         try:
#             # Decode the token
#             data = jwt.decode(token, app.config['JWT_SECRET_KEY'], algorithms=["HS256"])
#             current_user = data['username']
#         except IndexError
#             return jsonify({'error': 'Token has expired!'}), 401
#         except jwt.InvalidTokenError:
#             return jsonify({'error': 'Token is invalid!'}), 401
#         
#         return f(current_user, *args, **kwargs)
#     
#     return decorated

def ensure_dirs():
    # Use /tmp directory which is writable on App Engine
    os.makedirs('/tmp/results', exist_ok=True)
    # Ensure ML upload directory exists
    ml_uploads = Path(__file__).parent / 'ml' / 'api' / 'uploads'
    ml_uploads.mkdir(parents=True, exist_ok=True)

# @app.route('/api/login', methods=['POST'])
# def login():
#     try:
#         data = request.get_json()
#         username = data.get('username')
#         password = data.get('password')
#         
#         if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
#             # Generate JWT token with 1 hour expiration
#             token = jwt.encode({
#                 'username': username,
#                 'exp': datetime.utcnow() + timedelta(hours=1)
#             }, app.config['JWT_SECRET_KEY'], algorithm="HS256")
#             
#             return jsonify({
#                 'message': 'Login successful',
#                 'token': token,
#                 'username': username
#             })
#         else:
#             return jsonify({'error': 'Invalid credentials'}), 401
#     except Exception as e:
#         return jsonify({'error': f'Login error: {str(e)}'}), 500

# @app.route('/api/logout', methods=['POST'])
# def logout():
#     return jsonify({'message': 'Logged out successfully'})

# @app.route('/api/verify-token', methods=['POST'])
# @token_required
# def verify_token(current_user):
#     return jsonify({
#         'message': 'Token is valid',
#         'username': current_user
#     })

# XSS Scanner API endpoints
@app.route('/api/scan-xss', methods=['POST'])
# @token_required
def scan_xss():
    """XSS vulnerability scanning endpoint"""
    return api_scan_xss("anonymous")

@app.route('/api/generate-xss-report', methods=['POST'])
# @token_required
def generate_xss_report():
    """Generate Word report for XSS vulnerabilities"""
    return api_generate_xss_report("anonymous")

@app.route('/api/xss-sonarqube-export', methods=['POST'])
# @token_required
def xss_sonarqube_export():
    """Export XSS vulnerabilities in SonarQube format"""
    return api_xss_sonarqube_export("anonymous")

# SQL Injection Scanner API endpoints
@app.route('/api/scan-sql-injection', methods=['POST'])
# @token_required
def scan_sql_injection():
    """SQL injection vulnerability scanning endpoint"""
    return api_scan_sql_injection("anonymous")

@app.route('/api/generate-sql-injection-report', methods=['POST'])
# @token_required
def generate_sql_injection_report():
    """Generate Word report for SQL injection vulnerabilities"""
    return api_generate_sql_injection_report("anonymous")

@app.route('/api/sql-injection-sonarqube-export', methods=['POST'])
# @token_required
def sql_injection_sonarqube_export():
    """Export SQL injection vulnerabilities in SonarQube format"""
    return api_sql_injection_sonarqube_export("anonymous")

# Command Injection Scanner API endpoints
@app.route('/api/scan-command-injection', methods=['POST'])
# @token_required
def scan_command_injection():
    """Command injection vulnerability scanning endpoint"""
    return api_scan_command_injection("anonymous")

@app.route('/api/generate-command-injection-report', methods=['POST'])
# @token_required
def generate_command_injection_report():
    """Generate Word report for command injection vulnerabilities"""
    return api_generate_command_injection_report("anonymous")

@app.route('/api/command-injection-sonarqube-export', methods=['POST'])
# @token_required
def command_injection_sonarqube_export():
    """Export command injection vulnerabilities in SonarQube format"""
    return api_command_injection_sonarqube_export("anonymous")

# CSRF Scanner API endpoints
@app.route('/api/scan-csrf', methods=['POST'])
# @token_required
def scan_csrf():
    """CSRF vulnerability scanning endpoint"""
    return api_scan_csrf("anonymous")

@app.route('/api/generate-csrf-report', methods=['POST'])
# @token_required
def generate_csrf_report():
    """Generate Word report for CSRF vulnerabilities"""
    return api_generate_csrf_report("anonymous")

@app.route('/api/csrf-sonarqube-export', methods=['POST'])
# @token_required
def csrf_sonarqube_export():
    """Export CSRF vulnerabilities in SonarQube format"""
    return api_csrf_sonarqube_export("anonymous")

# Configuration endpoint
@app.route('/api/scanner-config', methods=['GET'])
def get_scanner_config():
    """Get scanner configuration"""
    try:
        config_path = os.path.join(os.path.dirname(__file__), 'scanner_config.json')
        with open(config_path, 'r') as f:
            config = json.load(f)
        return jsonify(config)
    except Exception as e:
        return jsonify({'error': f'Failed to load configuration: {str(e)}'}), 500

# ML report generation endpoint
@app.route('/api/generate-ml-report', methods=['POST'])
def generate_ml_report():
    """Generate a Word report for ML-based analysis including the visualization image."""
    try:
        data = request.get_json(force=True)
        upload_id = data.get('upload_id')
        image_url = data.get('image_url') or data.get('image')
        original_code = data.get('original_code') or data.get('code') or ''
        filename = data.get('filename') or data.get('file_name') or 'code.py'
        scanner_type = (data.get('scanner_type') or data.get('type') or 'ml').upper()

        if not upload_id or not image_url:
            return jsonify({'error': 'upload_id and image_url are required'}), 400

        # Expecting image_url in the form: /api/ml-output/<uid>/<filename>
        try:
            _, api_prefix, ml_output, uid, image_name = image_url.split('/', 4)
            if ml_output != 'ml-output' or uid != upload_id:
                raise ValueError('Mismatched upload id in image_url')
        except Exception:
            return jsonify({'error': 'image_url must be like /api/ml-output/<upload_id>/<image_file>.png'}), 400

        backend_root = Path(__file__).parent
        image_path = backend_root / 'ml' / 'api' / 'uploads' / upload_id / 'output' / image_name
        if not image_path.exists():
            return jsonify({'error': f'Visualization not found on server: {str(image_path)}'}), 404

        # Build the document
        doc = Document()
        # Use landscape Letter with reasonable margins and scale image to fit page width
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        title = doc.add_heading('ML-Based Security Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Visualization
        doc.add_heading('Visualization', level=1)
        try:
            # Fit image to available page width (about ~2x previous 6.5")
            available_width = section.page_width - section.left_margin - section.right_margin
            doc.add_picture(str(image_path), width=available_width)
        except Exception:
            # In case the image is extremely wide, add without width and let Word handle scaling
            doc.add_picture(str(image_path))

        # Save to temp file and return
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(tmp.name)
        tmp.close()

        download_name = f"ml_security_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=download_name,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': f'ML report generation error: {str(e)}'}), 500

# ML-based analysis endpoint
@app.route('/api/scan-ml', methods=['POST'])
def scan_ml():
    """Run machine learning based analysis using LSTM models (Atiqullah Ahmadzai’s project)."""
    try:
        data = request.get_json(force=True)
        vuln_type = (data.get('type') or '').lower()
        code = data.get('code', '')
        filename = data.get('filename') or 'code.py'

        if not code or not vuln_type:
            return jsonify({'error': 'type and code are required'}), 400

        # Map UI types to model modes
        mode_map = {
            'sql': 'sql',
            'xss': 'xss',
            'command': 'command_injection',
            'csrf': 'xsrf'
        }
        if vuln_type not in mode_map:
            return jsonify({'error': f'Unsupported type: {vuln_type}'}), 400
        mode = mode_map[vuln_type]

        # Prepare upload folder and file
        uid = uuid.uuid4().hex
        uploads_dir = Path(__file__).parent / 'ml' / 'api' / 'uploads' / uid
        uploads_dir.mkdir(parents=True, exist_ok=True)
        file_path = uploads_dir / filename
        file_path.write_text(code)

        # Choose python interpreter for ML
        backend_root = Path(__file__).parent
        ml_python_candidates = [
            backend_root / 'mlvenv' / 'bin' / 'python',
            Path(__file__).parent.parent / 'venv' / 'bin' / 'python',
        ]
        python_bin = None
        for candidate in ml_python_candidates:
            if candidate.exists():
                python_bin = str(candidate)
                break
        if python_bin is None:
            python_bin = 'python3'

        # Call the REAL ML pipeline (Atiqullah Ahmadzai's demonstrate.py)
        ml_api_cwd = backend_root / 'ml' / 'api'
        log_id = uuid.uuid4().hex[:8]
        output_dir = ml_api_cwd / 'uploads' / uid / 'output'
        output_dir.mkdir(parents=True, exist_ok=True)

        try:
            completed = subprocess.run(
                [python_bin, '../lib/demonstrate.py', mode, uid, filename, log_id],
                cwd=str(ml_api_cwd),
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=300
            )
        except subprocess.CalledProcessError as e:
            return jsonify({
                'status': 'error',
                'type': vuln_type,
                'mode': mode,
                'upload_id': uid,
                'filename': filename,
                'message': 'ML analysis failed',
                'stderr': e.stderr[-1000:],
                'stdout': e.stdout[-1000:]
            }), 500
        except subprocess.TimeoutExpired:
            return jsonify({
                'status': 'timeout',
                'type': vuln_type,
                'mode': mode,
                'upload_id': uid,
                'filename': filename,
                'message': 'ML analysis timed out'
            }), 504

        # The demonstrate.py writes: ../api/uploads/<uid>/output/<log_id>_<filename>_<mode>.png
        output_filename = f"{log_id}_{filename}_{mode}.png"
        expected_path = output_dir / output_filename
        if not expected_path.exists():
            # Try to find any generated file for debugging
            generated = sorted(output_dir.glob(f"*_{filename}_{mode}.png"))
            alt = generated[-1].name if generated else None
            return jsonify({
                'status': 'error',
                'type': vuln_type,
                'mode': mode,
                'upload_id': uid,
                'filename': filename,
                'message': 'ML image not found after run',
                'expected': output_filename,
                'found': alt
            }), 500

        image_url = f"/api/ml-output/{uid}/{output_filename}"
        return jsonify({
            'status': 'completed',
            'type': vuln_type,
            'mode': mode,
            'upload_id': uid,
            'filename': filename,
            'image_url': image_url
        })
    except Exception as e:
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500

@app.route('/api/ml-output/<string:uid>/<path:filename>', methods=['GET'])
def get_ml_output(uid: str, filename: str):
    """Serve generated ML visualization images from ml/api/uploads/<uid>/output."""
    try:
        backend_root = Path(__file__).parent
        file_path = backend_root / 'ml' / 'api' / 'uploads' / uid / 'output' / filename
        if not file_path.exists():
            return jsonify({'error': 'File not found'}), 404
        return send_file(str(file_path), mimetype='image/png')
    except Exception as e:
        return jsonify({'error': f'Failed to serve file: {str(e)}'}), 500

# Health check endpoint
@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'version': '1.0.0',
        'scanners': ['xss', 'sql_injection', 'command_injection', 'csrf']
    })

# Initialize directories on startup
ensure_dirs()

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    debug = os.environ.get('FLASK_ENV') != 'production'
    app.run(host='0.0.0.0', port=port, debug=debug)


# cd /Users/mohammen.almasi/thesis/06.27 && source venv/bin/activate && cd backend && python main.py
# cd frontend && npm start