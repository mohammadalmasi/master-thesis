from flask import request, jsonify, send_file
import requests
import re
import os
import ast
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
try:
    from sonarqube_security_standards import SecurityStandards, SQCategory, VulnerabilityProbability
    # Based on SonarQube's SecurityStandards.java
    SQL_INJECTION_VULNERABILITY = ("sql_injection", VulnerabilityProbability.HIGH)
    NOSQL_INJECTION_VULNERABILITY = ("nosql_injection", VulnerabilityProbability.HIGH)
except ImportError:
    # Fallback if sonarqube_security_standards is not available
    class VulnerabilityProbability:
        HIGH = "HIGH"
    SQL_INJECTION_VULNERABILITY = ("sql_injection", VulnerabilityProbability.HIGH)
    NOSQL_INJECTION_VULNERABILITY = ("nosql_injection", VulnerabilityProbability.HIGH)
    
from datetime import datetime
import json
import tempfile
# Maps to CWE-89, CWE-564, CWE-943
# Maps to OWASP A03:2021-Injection

class SQLInjectionVulnerability:
    def __init__(self, line_number, vulnerability_type, description, severity, code_snippet, remediation, confidence, file_path=None):
        self.line_number = line_number
        self.vulnerability_type = vulnerability_type
        self.description = description
        self.severity = severity
        self.code_snippet = code_snippet
        self.remediation = remediation
        self.confidence = confidence
        self.file_path = file_path or 'unknown'
        
    def to_dict(self):
        return {
            'line_number': self.line_number,
            'vulnerability_type': self.vulnerability_type,
            'description': self.description,
            'severity': self.severity,
            'code_snippet': self.code_snippet,
            'remediation': self.remediation,
            'confidence': self.confidence,
            'file_path': self.file_path,
            'cwe_references': ["89", "564", "943"],
            'owasp_references': ["A03:2021-Injection"],
            'rule_key': 'python:S2077'  # SQL injection rule key similar to SonarQube
        }

class SQLInjectionDetector:
    def __init__(self, debug=False):
        self.vulnerabilities = []
        self.debug = debug
        
    def scan_file(self, filename):
        """Scan a Python file for SQL injection vulnerabilities"""
        self.vulnerabilities = []
        
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                code = f.read()
        except UnicodeDecodeError:
            with open(filename, 'r', encoding='latin-1') as f:
                code = f.read()
        
        # Scan for SQL injection patterns
        self._scan_for_patterns(code, filename)
        self._scan_with_ast(code, filename)
        
        # Remove duplicates
        self.vulnerabilities = self._deduplicate_vulnerabilities(self.vulnerabilities)
        
        return self.vulnerabilities
    
    def _deduplicate_vulnerabilities(self, vulnerabilities):
        """Remove duplicate vulnerabilities based on line number and similarity"""
        if not vulnerabilities:
            return vulnerabilities
        
        # Group vulnerabilities by line number
        by_line = {}
        for vuln in vulnerabilities:
            line_key = vuln.line_number
            if line_key not in by_line:
                by_line[line_key] = []
            by_line[line_key].append(vuln)
        
        # For each line, keep only the best vulnerability
        deduplicated = []
        for line_num, line_vulns in by_line.items():
            if len(line_vulns) == 1:
                deduplicated.extend(line_vulns)
            else:
                # Multiple vulnerabilities on the same line - keep the best one
                # Sort by confidence (highest first), then by description length (longest first)
                best_vuln = max(line_vulns, key=lambda v: (v.confidence, len(v.description)))
                deduplicated.append(best_vuln)
        
        return deduplicated
    
    def _scan_for_patterns(self, code, filename):
        """Scan for SQL injection vulnerability patterns using regex"""
        lines = code.split('\n')
        
        # Pre-process code to handle line continuations for multi-line patterns
        processed_code = self._preprocess_multiline_code(code)
        
        # Also scan the original line-by-line for single-line patterns
        self._scan_lines_for_patterns(lines, filename)
        
        # Scan processed code for multi-line patterns
        self._scan_multiline_patterns(processed_code, filename)
    
    def _preprocess_multiline_code(self, code):
        """Pre-process code to handle line continuations and create single-line equivalents"""
        lines = code.split('\n')
        processed_lines = []
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Check if line ends with backslash (continuation)
            if line.endswith('\\'):
                # Start building a multi-line statement
                multiline_parts = [line[:-1].rstrip()]  # Remove backslash
                original_line_num = i + 1
                i += 1
                
                # Continue collecting lines until we find one without backslash
                while i < len(lines):
                    next_line = lines[i].strip()
                    if next_line.endswith('\\'):
                        multiline_parts.append(next_line[:-1].rstrip())
                        i += 1
                    else:
                        multiline_parts.append(next_line)
                        break
                
                # Join the multiline statement
                multiline_statement = ' '.join(multiline_parts)
                processed_lines.append((original_line_num, multiline_statement))
            else:
                processed_lines.append((i + 1, line))
            
            i += 1
        
        return processed_lines
    
    def _scan_multiline_patterns(self, processed_lines, filename):
        """Scan for multi-line SQL injection patterns"""
        multiline_patterns = [
            # Multi-line string concatenation with SQL keywords
            {
                'pattern': r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE)[^\'\"]*[\'\"]\s*\+.*?\+\s*\w+)',
                'description': 'Multi-line SQL query with string concatenation - potential SQL injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries instead of string concatenation'
            },
            # Multi-line concatenation with user input variables
            {
                'pattern': r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE)[^\'\"]*[\'\"]\s*\+.*?\+\s*(?:search_term|user_id|category|username|email|input|param|data)\s*\+)',
                'description': 'Multi-line SQL concatenation with user input variables - SQL injection risk',
                'severity': 'high',
                'confidence': 0.85,
                'remediation': 'Use parameterized queries to prevent SQL injection'
            },
            # Multi-line WHERE clause concatenation
            {
                'pattern': r'([\'\"]\s*WHERE[^\'\"]*[\'\"]\s*\+.*?\+\s*\w+)',
                'description': 'Multi-line WHERE clause with concatenation - SQL injection risk',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use parameterized WHERE clauses'
            },
            # Multi-line LIKE clause concatenation
            {
                'pattern': r'([\'\"]\s*(?:LIKE|=)[^\'\"]*[\'\"]\s*\+.*?\+\s*\w+)',
                'description': 'Multi-line LIKE/comparison with concatenation - SQL injection risk',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use parameterized queries for LIKE clauses'
            }
        ]
        
        for line_num, line_content in processed_lines:
            for pattern_info in multiline_patterns:
                matches = re.finditer(pattern_info['pattern'], line_content, re.IGNORECASE | re.DOTALL)
                for match in matches:
                    vulnerability = SQLInjectionVulnerability(
                        line_number=line_num,
                        vulnerability_type='sql_injection',
                        description=pattern_info['description'],
                        severity=pattern_info['severity'],
                        code_snippet=line_content.strip(),
                        remediation=pattern_info['remediation'],
                        confidence=pattern_info['confidence'],
                        file_path=filename
                    )
                    self.vulnerabilities.append(vulnerability)
    
    def _scan_lines_for_patterns(self, lines, filename):
        """Scan individual lines for single-line SQL injection patterns"""
        
        # SQL injection vulnerability patterns
        patterns = [
            # MEDIUM SEVERITY PATTERNS (Check these first for specificity)
            # ORDER BY patterns - both quoted and unquoted
            {
                'pattern': r'([\'"]ORDER\s+BY[^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'ORDER BY clause with string concatenation - SQL injection risk',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Validate column names or use parameterized queries'
            },
            {
                'pattern': r'([\'"]ORDER\s+BY\s*[\'"]\s*\+\s*\w+)',
                'description': 'ORDER BY clause with string concatenation - SQL injection risk',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Validate column names or use parameterized queries'
            },
            # LIMIT patterns - both quoted and unquoted
            {
                'pattern': r'([\'"]LIMIT[^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'LIMIT clause with string concatenation - SQL injection risk',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Use parameterized LIMIT clauses'
            },
            {
                'pattern': r'([\'"]LIMIT\s*[\'"]\s*\+\s*\w+)',
                'description': 'LIMIT clause with string concatenation - SQL injection risk',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Use parameterized LIMIT clauses'
            },
            # SQL comments injection
            {
                'pattern': r'(\w+\s*\+\s*[\'"].*--.*[\'"])',
                'description': 'SQL comment injection pattern detected',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Remove or escape SQL comment characters'
            },
            # Blind SQL injection patterns
            {
                'pattern': r'(\w+\s*\+\s*[\'"].*(?:AND|OR)\s+\d+\s*=\s*\d+)',
                'description': 'Potential blind SQL injection pattern',
                'severity': 'medium',
                'confidence': 0.6,
                'remediation': 'Use parameterized queries and input validation'
            },
            # HIGH SEVERITY PATTERNS (More general)
            # General string concatenation with SQL keywords
            {
                'pattern': r'([\'\"]\s*.*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*[\'\"]\s*\+\s*\w+)',
                'description': 'SQL query with string concatenation - potential SQL injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries or prepared statements'
            },
            # String concatenation in the middle of SQL (excluding ORDER and LIMIT which are handled above)
            {
                'pattern': r'(\w+\s*\+\s*[\'"].*(?:SELECT|INSERT|UPDATE|DELETE|WHERE|FROM|JOIN|UNION).*[\'"])',
                'description': 'String concatenation with SQL keywords - potential SQL injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries or prepared statements'
            },
            # Direct string concatenation in SQL queries - more flexible
            {
                'pattern': r'((?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE)\s+[^"\']*[\'"]\s*\+\s*\w+)',
                'description': 'SQL query with string concatenation - potential SQL injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries or prepared statements'
            },
            {
                'pattern': r'(\w+\s*\+\s*[\'"](?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE))',
                'description': 'String concatenation with SQL keywords - potential SQL injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries or prepared statements'
            },
            # F-string patterns with SQL
            {
                'pattern': r'(f[\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?\{[^}]*\}.*[\'"])',
                'description': 'F-string with SQL query and user input - potential SQL injection',
                'severity': 'high',
                'confidence': 0.95,
                'remediation': 'Use parameterized queries instead of f-strings for SQL'
            },
            # Format string patterns
            {
                'pattern': r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?\{\}.*?[\'\"]\s*\.format\s*\([^)]*\))',
                'description': 'Format string with SQL query - potential SQL injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries instead of format strings'
            },
            # % string formatting
            {
                'pattern': r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?%[sd].*?[\'\"]\s*%\s*[^;]+)',
                'description': 'String formatting with SQL query - potential SQL injection',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries instead of % formatting'
            },
            # cursor.execute with concatenation
            {
                'pattern': r'(cursor\.execute\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'cursor.execute() with string concatenation - SQL injection risk',
                'severity': 'high',
                'confidence': 0.95,
                'remediation': 'Use parameterized queries with cursor.execute(sql, params)'
            },
            # execute with concatenation
            {
                'pattern': r'(\.execute\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Database execute() with string concatenation - SQL injection risk',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries with execute(sql, params)'
            },
            # Raw SQL execution patterns
            {
                'pattern': r'(cursor\.execute\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
                'description': 'cursor.execute() with f-string - SQL injection risk',
                'severity': 'high',
                'confidence': 0.95,
                'remediation': 'Use parameterized queries instead of f-strings'
            },
            # Django ORM raw queries
            {
                'pattern': r'(\.raw\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Django raw() query with concatenation - SQL injection risk',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized raw queries or ORM methods'
            },
            {
                'pattern': r'(\.raw\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
                'description': 'Django raw() query with f-string - SQL injection risk',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized raw queries instead of f-strings'
            },
            # SQLAlchemy text() with concatenation
            {
                'pattern': r'(text\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'SQLAlchemy text() with concatenation - SQL injection risk',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries with text(sql, params)'
            },
            {
                'pattern': r'(text\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
                'description': 'SQLAlchemy text() with f-string - SQL injection risk',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries instead of f-strings'
            },
            # More specific SQLAlchemy text() patterns
            {
                'pattern': r'(\w+\s*=\s*text\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
                'description': 'SQLAlchemy text() assignment with f-string - SQL injection risk',
                'severity': 'high',
                'confidence': 0.95,
                'remediation': 'Use parameterized queries instead of f-strings: text("SELECT * FROM users WHERE id = :user_id", user_id=user_id)'
            },
            {
                'pattern': r'(return\s+text\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
                'description': 'SQLAlchemy text() return with f-string - SQL injection risk',
                'severity': 'high',
                'confidence': 0.95,
                'remediation': 'Use parameterized queries instead of f-strings'
            },
            {
                'pattern': r'(text\s*\(\s*[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
                'description': 'SQLAlchemy text() with string formatting - SQL injection risk',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries instead of string formatting'
            },
            # WHERE clause patterns
            {
                'pattern': r'([\'"]WHERE\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'WHERE clause with string concatenation - SQL injection risk',
                'severity': 'high',
                'confidence': 0.85,
                'remediation': 'Use parameterized WHERE clauses'
            },
            # LIKE patterns
            {
                'pattern': r'([\'"]LIKE\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'LIKE clause with string concatenation - SQL injection risk',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use parameterized LIKE clauses'
            },
            # Request parameter usage in SQL
            {
                'pattern': r'(request\.args\.get\([^)]*\)[^;]*(?:SELECT|INSERT|UPDATE|DELETE|WHERE|LIKE))',
                'description': 'Request parameters used directly in SQL context',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Sanitize and validate request parameters before using in SQL'
            },
            {
                'pattern': r'(request\.form\.get\([^)]*\)[^;]*(?:SELECT|INSERT|UPDATE|DELETE|WHERE|LIKE))',
                'description': 'Form data used directly in SQL context',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Sanitize and validate form data before using in SQL'
            },
            # Request parameter usage with bracket notation
            {
                'pattern': r'(request\.args\[[^]]*\])',
                'description': 'Request parameter accessed directly - potential injection risk',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Validate and sanitize request parameters before use'
            },
            {
                'pattern': r'(request\.form\[[^]]*\])',
                'description': 'Form data accessed directly - potential injection risk',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Validate and sanitize form data before use'
            },
            # Unsafe SQL construction
            {
                'pattern': r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE)[^\'\"]*[\'"]\s*%\s*\([^)]*\))',
                'description': 'SQL query with dictionary formatting - potential injection',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use parameterized queries instead of string formatting'
            },
            # Dynamic table/column names
            {
                'pattern': r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE)[^\'\"]*FROM\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Dynamic table name construction - SQL injection risk',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Validate table names against whitelist or use ORM methods'
            },
            # UNION injection patterns
            {
                'pattern': r'(\w+\s*\+\s*[\'"].*UNION.*SELECT)',
                'description': 'String concatenation with UNION SELECT - SQL injection risk',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries and input validation'
            },
            # Raw SQL injection patterns (without Python syntax)
            {
                'pattern': r'(SELECT\s+[^{]*\{[^}]*(?:user_input|username|input|param|data|request|form|args)\}[^}]*UNION\s+SELECT)',
                'description': 'Raw SQL with placeholder and UNION SELECT - SQL injection risk',
                'severity': 'high',
                'confidence': 0.95,
                'remediation': 'Use parameterized queries instead of placeholder injection'
            },
            {
                'pattern': r'((?:SELECT|INSERT|UPDATE|DELETE)\s+[^{]*\{[^}]*(?:user_input|username|input|param|data|request|form|args)\})',
                'description': 'Raw SQL with user input placeholders - SQL injection risk',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized queries instead of placeholder injection'
            },
            {
                'pattern': r'((?:SELECT|INSERT|UPDATE|DELETE)[^U]*UNION\s+SELECT)',
                'description': 'Raw SQL with UNION SELECT - potential SQL injection',
                'severity': 'high',
                'confidence': 0.85,
                'remediation': 'Validate and sanitize all user inputs to prevent UNION-based attacks'
            },
            {
                'pattern': r'(WHERE\s+[^{]*\{[^}]*(?:user_input|username|input|param|data|request|form|args)\})',
                'description': 'Raw SQL WHERE clause with user input placeholders - SQL injection risk',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Use parameterized WHERE clauses instead of placeholder injection'
            },
            # Stored procedure calls
            {
                'pattern': r'(CALL\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'Stored procedure call with concatenation - SQL injection risk',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use parameterized stored procedure calls'
            },
            # MySQL specific patterns
            {
                'pattern': r'(LOAD_FILE\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'MySQL LOAD_FILE with concatenation - SQL injection risk',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Validate file paths and use parameterized queries'
            },
            # LOW SEVERITY PATTERNS
            # Simple string concatenation without obvious SQL keywords
            {
                'pattern': r'(\w+_\w+\s*\+\s*\w+\s*$)',
                'description': 'Simple string concatenation - potential low-risk SQL injection',
                'severity': 'low',
                'confidence': 0.4,
                'remediation': 'Validate input and consider parameterized queries'
            },
            # Variable building patterns
            {
                'pattern': r'([\'"][\w_]+[\'\"]\s*\+\s*\w+\s*$)',
                'description': 'Basic string building - potential low-risk injection',
                'severity': 'low',
                'confidence': 0.3,
                'remediation': 'Ensure proper input validation'
            },
            # Column/table name concatenation
            {
                'pattern': r'(\w+\s*=\s*[\'"][\w_]+[\'\"]\s*\+\s*\w+)',
                'description': 'Dynamic column/table name construction - low risk injection',
                'severity': 'low',
                'confidence': 0.4,
                'remediation': 'Use whitelist validation for column/table names'
            },
            
            # ============================================================================
            # NOSQL INJECTION PATTERNS (MongoDB, etc.)
            # ============================================================================
            
            # MongoDB collection.find() with user input
            {
                'pattern': r'(collection\.find\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input)[^}]*\})',
                'description': 'MongoDB find() with user input - NoSQL injection risk',
                'severity': 'high',
                'confidence': 0.85,
                'remediation': 'Validate and sanitize user input before MongoDB queries'
            },
            {
                'pattern': r'(\.find\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input)[^}]*\})',
                'description': 'MongoDB find() with user input - NoSQL injection risk',
                'severity': 'high',
                'confidence': 0.85,
                'remediation': 'Validate and sanitize user input before MongoDB queries'
            },
            
            # MongoDB db.eval() - extremely dangerous
            {
                'pattern': r'(db\.eval\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
                'description': 'MongoDB eval() with f-string - critical NoSQL injection risk',
                'severity': 'high',
                'confidence': 0.95,
                'remediation': 'Never use db.eval() with user input - consider aggregation pipeline instead'
            },
            {
                'pattern': r'(db\.eval\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
                'description': 'MongoDB eval() with string concatenation - critical NoSQL injection risk',
                'severity': 'high',
                'confidence': 0.95,
                'remediation': 'Never use db.eval() with user input - consider aggregation pipeline instead'
            },
            {
                'pattern': r'(db\.eval\s*\([^)]*(?:request\.|username|user_input|\w+_input))',
                'description': 'MongoDB eval() with user input - critical NoSQL injection risk',
                'severity': 'high',
                'confidence': 0.9,
                'remediation': 'Never use db.eval() with user input - extremely dangerous'
            },
            
            # MongoDB other methods with user input
            {
                'pattern': r'(\.(?:find_one|update|delete|remove|insert)\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input))',
                'description': 'MongoDB operation with user input - NoSQL injection risk',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Validate and sanitize user input before MongoDB operations'
            },
            
            # Direct user input in MongoDB query dictionary
            {
                'pattern': r'(\{[^}]*[\'"]:\s*request\.(?:form|args|json)\[[^]]*\][^}]*\})',
                'description': 'Direct request parameter in MongoDB query - NoSQL injection risk',
                'severity': 'high',
                'confidence': 0.85,
                'remediation': 'Validate request parameters before using in MongoDB queries'
            },
            
            # MongoDB query with format string
            {
                'pattern': r'(\{[^}]*[\'"]:\s*[\'"][^\'\"]*\{\}[^\'\"]*[\'\"]\s*\.format\s*\([^)]*\))',
                'description': 'MongoDB query with format string - NoSQL injection risk',
                'severity': 'high',
                'confidence': 0.8,
                'remediation': 'Use proper MongoDB query parameters instead of format strings'
            },
            
            # MongoDB aggregation pipeline with user input
            {
                'pattern': r'(aggregate\s*\(\s*\[[^]]*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input))',
                'description': 'MongoDB aggregation with user input - NoSQL injection risk',
                'severity': 'medium',
                'confidence': 0.7,
                'remediation': 'Validate user input in aggregation pipelines'
            },
            
            # Generic NoSQL client patterns
            {
                'pattern': r'(pymongo\.MongoClient[^;]*find\s*\([^)]*(?:request\.|username|user_input))',
                'description': 'PyMongo query with user input - NoSQL injection risk',
                'severity': 'high',
                'confidence': 0.75,
                'remediation': 'Sanitize user input before MongoDB queries'
            }
        ]
        
        for line_num, line in enumerate(lines, 1):
            for pattern_info in patterns:
                matches = re.finditer(pattern_info['pattern'], line, re.IGNORECASE)
                for match in matches:
                    vulnerability = SQLInjectionVulnerability(
                        line_number=line_num,
                        vulnerability_type='sql_injection',
                        description=pattern_info['description'],
                        severity=pattern_info['severity'],
                        code_snippet=line.strip(),
                        remediation=pattern_info['remediation'],
                        confidence=pattern_info['confidence'],
                        file_path=filename
                    )
                    self.vulnerabilities.append(vulnerability)
    
    def _scan_with_ast(self, code, filename):
        """Scan for SQL injection vulnerabilities using AST analysis"""
        try:
            tree = ast.parse(code)
            visitor = SQLInjectionASTVisitor(filename)
            visitor.visit(tree)
            self.vulnerabilities.extend(visitor.vulnerabilities)
        except SyntaxError:
            # If code has syntax errors, skip AST analysis
            pass

class SQLInjectionASTVisitor(ast.NodeVisitor):
    def __init__(self, filename):
        self.filename = filename
        self.vulnerabilities = []
        self.sql_keywords = {
            'SELECT', 'INSERT', 'UPDATE', 'DELETE', 'CREATE', 'DROP', 'ALTER',
            'EXEC', 'EXECUTE', 'WHERE', 'FROM', 'JOIN', 'UNION', 'ORDER', 'GROUP'
        }
        
        # NoSQL method names that can be vulnerable
        self.nosql_methods = {
            'find', 'find_one', 'update', 'delete', 'remove', 'insert', 'eval',
            'aggregate', 'count', 'distinct', 'find_and_modify'
        }
    
    def visit_Call(self, node):
        """Visit function calls to detect SQL injection patterns"""
        # Check for database execute methods
        if isinstance(node.func, ast.Attribute):
            if (hasattr(node.func, 'attr') and 
                node.func.attr in ['execute', 'executemany', 'raw']):
                
                # Check if arguments contain SQL injection patterns
                for arg in node.args:
                    if self._contains_sql_injection_pattern(arg):
                        vulnerability = SQLInjectionVulnerability(
                            line_number=getattr(node, 'lineno', 0),
                            vulnerability_type='sql_injection',
                            description=f'Potentially unsafe {node.func.attr}() call with dynamic SQL',
                            severity='high',
                            code_snippet=self._get_code_snippet(node),
                            remediation='Use parameterized queries instead of string concatenation',
                            confidence=0.8,
                            file_path=self.filename
                        )
                        self.vulnerabilities.append(vulnerability)
            
            # Check for NoSQL methods (MongoDB, etc.)
            elif (hasattr(node.func, 'attr') and 
                  node.func.attr in self.nosql_methods):
                
                # Check if arguments contain NoSQL injection patterns
                for arg in node.args:
                    if self._contains_nosql_injection_pattern(arg):
                        # Special handling for db.eval() - extremely dangerous
                        if node.func.attr == 'eval':
                            vulnerability = SQLInjectionVulnerability(
                                line_number=getattr(node, 'lineno', 0),
                                vulnerability_type='nosql_injection',
                                description=f'Critical MongoDB eval() with user input - allows arbitrary code execution',
                                severity='high',
                                code_snippet=self._get_code_snippet(node),
                                remediation='Never use db.eval() with user input - use aggregation pipeline instead',
                                confidence=0.95,
                                file_path=self.filename
                            )
                        else:
                            vulnerability = SQLInjectionVulnerability(
                                line_number=getattr(node, 'lineno', 0),
                                vulnerability_type='nosql_injection',
                                description=f'MongoDB {node.func.attr}() with user input - NoSQL injection risk',
                                severity='high',
                                code_snippet=self._get_code_snippet(node),
                                remediation='Validate and sanitize user input before NoSQL queries',
                                confidence=0.8,
                                file_path=self.filename
                            )
                        self.vulnerabilities.append(vulnerability)
        
        # Check for SQLAlchemy text() function calls
        elif isinstance(node.func, ast.Name) and hasattr(node.func, 'id') and node.func.id == 'text':
            # Check if arguments contain SQL injection patterns
            for arg in node.args:
                if self._contains_sql_injection_pattern(arg):
                    vulnerability = SQLInjectionVulnerability(
                        line_number=getattr(node, 'lineno', 0),
                        vulnerability_type='sql_injection',
                        description='SQLAlchemy text() function with dynamic SQL - SQL injection risk',
                        severity='high',
                        code_snippet=self._get_code_snippet(node),
                        remediation='Use parameterized queries: text("SELECT * FROM users WHERE id = :user_id", user_id=user_id)',
                        confidence=0.9,
                        file_path=self.filename
                    )
                    self.vulnerabilities.append(vulnerability)
        
        self.generic_visit(node)
    
    def visit_BinOp(self, node):
        """Visit binary operations for string concatenation in SQL context"""
        if isinstance(node.op, ast.Add):
            # Check for SQL-like patterns in string concatenation
            left_str = self._extract_string_value(node.left)
            right_str = self._extract_string_value(node.right)
            
            if self._contains_sql_keywords(left_str) or self._contains_sql_keywords(right_str):
                if self._contains_user_input(node.right) or self._contains_user_input(node.left):
                    vulnerability = SQLInjectionVulnerability(
                        line_number=getattr(node, 'lineno', 0),
                        vulnerability_type='sql_injection',
                        description='String concatenation with SQL keywords and potential user input',
                        severity='high',
                        code_snippet=self._get_code_snippet(node),
                        remediation='Use parameterized queries instead of string concatenation',
                        confidence=0.8,
                        file_path=self.filename
                    )
                    self.vulnerabilities.append(vulnerability)
        
        self.generic_visit(node)
    
    def visit_JoinedStr(self, node):
        """Visit f-strings (JoinedStr nodes) for SQL injection patterns"""
        # Check if f-string contains SQL keywords
        sql_pattern = False
        has_user_input = False
        
        for value in node.values:
            if isinstance(value, ast.Str):
                if self._contains_sql_keywords(value.s):
                    sql_pattern = True
            elif isinstance(value, ast.Constant) and isinstance(value.value, str):
                if self._contains_sql_keywords(value.value):
                    sql_pattern = True
            elif isinstance(value, ast.FormattedValue):
                # Check if the formatted value contains user input
                if self._contains_user_input(value.value):
                    has_user_input = True
        
        if sql_pattern and has_user_input:
            vulnerability = SQLInjectionVulnerability(
                line_number=getattr(node, 'lineno', 0),
                vulnerability_type='sql_injection',
                description='F-string with SQL keywords and user input can lead to SQL injection',
                severity='high',
                code_snippet=self._get_code_snippet(node),
                remediation='Use parameterized queries instead of f-strings for SQL',
                confidence=0.95,
                file_path=self.filename
            )
            self.vulnerabilities.append(vulnerability)
        
        self.generic_visit(node)
    
    def _contains_sql_injection_pattern(self, node):
        """Check if node contains SQL injection patterns"""
        if isinstance(node, ast.BinOp) and isinstance(node.op, ast.Add):
            # String concatenation pattern
            left_str = self._extract_string_value(node.left)
            right_str = self._extract_string_value(node.right)
            
            if (self._contains_sql_keywords(left_str) or 
                self._contains_sql_keywords(right_str)):
                return True
        
        elif isinstance(node, ast.JoinedStr):
            # F-string pattern
            for value in node.values:
                if isinstance(value, ast.Str) and self._contains_sql_keywords(value.s):
                    return True
                elif (isinstance(value, ast.Constant) and 
                      isinstance(value.value, str) and 
                      self._contains_sql_keywords(value.value)):
                    return True
        
        return False
    
    def _contains_nosql_injection_pattern(self, node):
        """Check if node contains NoSQL injection patterns"""
        # Check for dictionary with user input values
        if isinstance(node, ast.Dict):
            for value in node.values:
                if self._contains_user_input(value):
                    return True
        
        # Check for f-strings with user input
        elif isinstance(node, ast.JoinedStr):
            for value in node.values:
                if isinstance(value, ast.FormattedValue):
                    if self._contains_user_input(value.value):
                        return True
        
        # Check for string concatenation with user input
        elif isinstance(node, ast.BinOp) and isinstance(node.op, ast.Add):
            if self._contains_user_input(node.left) or self._contains_user_input(node.right):
                return True
        
        # Check for direct user input
        elif self._contains_user_input(node):
            return True
            
        return False
    
    def _contains_sql_keywords(self, text):
        """Check if text contains SQL keywords"""
        if not text:
            return False
        
        text_upper = text.upper()
        return any(keyword in text_upper for keyword in self.sql_keywords)
    
    def _contains_user_input(self, node):
        """Check if node contains patterns that might be user input"""
        if isinstance(node, ast.Attribute):
            if (hasattr(node, 'attr') and 
                node.attr in ['args', 'form', 'cookies', 'headers', 'json', 'data']):
                return True
            # Check for request.args patterns
            if (isinstance(node.value, ast.Attribute) and
                hasattr(node.value, 'attr') and
                node.value.attr in ['args', 'form', 'cookies', 'headers', 'json'] and
                isinstance(node.value.value, ast.Name) and
                node.value.value.id == 'request'):
                return True
        elif isinstance(node, ast.Call):
            if (isinstance(node.func, ast.Attribute) and
                hasattr(node.func, 'attr') and
                node.func.attr in ['get', 'getlist']):
                return True
        elif isinstance(node, ast.Name):
            # Check for variables that might contain user input
            # Expanded list to include common variable names that might contain user input
            user_input_vars = [
                'input', 'data', 'param', 'value', 'user_input', 'username', 'password',
                'name', 'email', 'id', 'query', 'search', 'filter', 'sort', 'order',
                'limit', 'offset', 'page', 'user', 'login', 'auth', 'token', 'key',
                'text', 'content', 'message', 'comment', 'description', 'title',
                'field', 'column', 'table', 'where', 'condition', 'criteria',
                # NoSQL specific variables
                'user_query', 'category', 'criteria'
            ]
            if hasattr(node, 'id') and node.id in user_input_vars:
                return True
        return False
    
    def _extract_string_value(self, node):
        """Extract string value from AST node if it's a string literal"""
        if isinstance(node, ast.Str):
            return node.s
        elif isinstance(node, ast.Constant) and isinstance(node.value, str):
            return node.value
        return None
    
    def _get_code_snippet(self, node):
        """Get a string representation of the AST node"""
        try:
            return ast.unparse(node)
        except AttributeError:
            # Fallback for Python < 3.9
            return f'Line {getattr(node, "lineno", "unknown")}'

def highlight_sql_injection_vulnerabilities(code, vulnerabilities=None):
    """Highlight SQL injection vulnerability patterns in code with severity-based colors"""
    if not vulnerabilities:
        return code
    
    # Split code into lines for line-by-line highlighting
    lines = code.split('\n')
    highlighted_lines = []
    
    # Group vulnerabilities by line number
    vuln_by_line = {}
    multiline_vulns = []
    
    for vuln in vulnerabilities:
        line_num = vuln.line_number
        # Check if this is a multi-line vulnerability
        if ('Multi-line' in vuln.description or 
            '\\' in vuln.code_snippet or 
            len(vuln.code_snippet.split()) > 20):  # Heuristic for multi-line
            multiline_vulns.append(vuln)
        else:
            if line_num not in vuln_by_line:
                vuln_by_line[line_num] = []
            vuln_by_line[line_num].append(vuln)
    
    # Handle multi-line vulnerabilities by marking all related lines
    for vuln in multiline_vulns:
        start_line = vuln.line_number
        # Find the range of lines that contain this vulnerability
        end_line = start_line
        
        # Look for line continuation patterns starting from the vulnerability line
        for i in range(start_line - 1, len(lines)):
            if i >= 0 and lines[i].rstrip().endswith('\\'):
                end_line = i + 2  # Continue to next line
            elif i >= 0 and ('+' in lines[i] or 'query' in lines[i]):
                end_line = max(end_line, i + 1)
            else:
                break
        
        # Mark all lines in the range
        for line_num in range(start_line, min(end_line + 1, len(lines) + 1)):
            if line_num not in vuln_by_line:
                vuln_by_line[line_num] = []
            vuln_by_line[line_num].append(vuln)
    
    # Process each line
    for line_num, line in enumerate(lines, 1):
        if line_num in vuln_by_line:
            # Find the highest severity for this line
            line_vulns = vuln_by_line[line_num]
            severities = [v.severity for v in line_vulns]
            
            # Determine CSS class based on highest severity
            if 'high' in severities:
                css_class = 'sql-injection-vuln-high'
            elif 'medium' in severities:
                css_class = 'sql-injection-vuln-medium'
            else:
                css_class = 'sql-injection-vuln-low'
            
            # For multi-line vulnerabilities, highlight the entire line
            # For single-line vulnerabilities, try to highlight specific patterns
            highlighted_line = line
            has_multiline_vuln = any('Multi-line' in v.description for v in line_vulns)
            
            if has_multiline_vuln:
                # For multi-line vulnerabilities, highlight the entire line
                highlighted_line = f'<span class="{css_class}">{line}</span>'
            else:
                # For single-line vulnerabilities, use pattern-based highlighting
                patterns = _create_highlighting_patterns(line_vulns)
                
                # Track if we've already highlighted this line to prevent overlapping spans
                already_highlighted = False
                
                for pattern in patterns:
                    try:
                        # Only apply highlighting if line hasn't been highlighted yet
                        if not already_highlighted:
                            new_line = re.sub(f'({pattern})', 
                                            lambda m: f'<span class="{css_class}">{m.group(0)}</span>', 
                                            highlighted_line, flags=re.IGNORECASE)
                            # Check if highlighting was applied
                            if new_line != highlighted_line:
                                highlighted_line = new_line
                                already_highlighted = True
                                break
                    except re.error:
                        # If regex fails, highlight the whole line
                        highlighted_line = f'<span class="{css_class}">{line}</span>'
                        break
            
            highlighted_lines.append(highlighted_line)
        else:
            highlighted_lines.append(line)
    
    return '\n'.join(highlighted_lines)

def _create_highlighting_patterns(line_vulns):
    """Create highlighting patterns for single-line vulnerabilities"""
    patterns = []
    
    # Common SQL injection patterns for highlighting
    common_patterns = [
        # High severity patterns
        r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?\{\}.*?[\'\"]\s*\.format\s*\([^)]*\))',
        r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?%[sd].*?[\'\"]\s*%\s*[^;]+)',
        r'(f[\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?\{[^}]*\}.*?[\'"])',
        r'([\'\"]\s*.*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*[\'\"]\s*\+\s*\w+)',
        r'(\w+\s*\+\s*[\'"].*(?:SELECT|INSERT|UPDATE|DELETE|WHERE|FROM|JOIN|UNION).*[\'"])',
        r'(cursor\.execute\s*\([^)]*\))',
        r'(\.execute\s*\([^)]*\))',
        r'(text\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(text\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'([\'"]WHERE\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
        r'([\'"]LIKE\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
        
        # Medium severity patterns - ORDER BY and LIMIT
        r'([\'"]ORDER\s+BY[^\'\"]*[\'"]\s*\+\s*\w+)',
        r'([\'"]ORDER\s+BY\s*[\'"]\s*\+\s*\w+)',
        r'([\'"]LIMIT[^\'\"]*[\'"]\s*\+\s*\w+)',
        r'([\'"]LIMIT\s*[\'"]\s*\+\s*\w+)',
        
        # Medium severity patterns - SQL comments and blind injection
        r'(\w+\s*\+\s*[\'"].*--.*[\'"])',
        r'(\w+\s*\+\s*[\'"].*(?:AND|OR)\s+\d+\s*=\s*\d+)',
        
        # Medium severity patterns - Request parameters
        r'(request\.(?:form|args)\[[^]]*\])',
        r'(request\.args\.get\([^)]*\))',
        r'(request\.form\.get\([^)]*\))',
        
        # Low severity patterns - Simple string concatenation
        r'(\w+\s*\+\s*\w+\s*$)',
        r'([\'"][\w_]+[\'\"]\s*\+\s*\w+)',
        r'(\w+\s*=\s*[\'"][\w_]+[\'\"]\s*\+\s*\w+)',
        
        # NoSQL injection patterns
        r'(collection\.find\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input)[^}]*\})',
        r'(\.find\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input)[^}]*\})',
        r'(db\.eval\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'(db\.eval\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(db\.eval\s*\([^)]*(?:request\.|username|user_input|\w+_input))',
        r'(\.(?:find_one|update|delete|remove|insert)\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input))',
        r'(\{[^}]*[\'"]:\s*request\.(?:form|args|json)\[[^]]*\][^}]*\})',
        r'(aggregate\s*\(\s*\[[^]]*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input))'
    ]
    
    # Add vulnerability-specific patterns based on the actual vulnerabilities found
    for vuln in line_vulns:
        snippet = vuln.code_snippet.strip()
        if snippet and len(snippet) < 200:  # Only for reasonable-length snippets
            # Create a flexible pattern from the snippet
            if '.format(' in snippet:
                patterns.append(r'[\'\"]\s*[^\'\"]*\{\}[^\'\"]*[\'\"]\s*\.format\s*\([^)]*\)')
            elif ' % ' in snippet and ('SELECT' in snippet.upper() or 'UPDATE' in snippet.upper()):
                patterns.append(r'[\'\"]\s*[^\'\"]*%[sd][^\'\"]*[\'\"]\s*%\s*[^;]+')
            elif 'request.' in snippet:
                patterns.append(r'request\.\w+\[[^]]*\]')
            elif 'ORDER BY' in snippet.upper() and ' + ' in snippet:
                patterns.append(r'[\'"]ORDER\s+BY[^\'\"]*[\'"]\s*\+\s*\w+')
            elif 'LIMIT' in snippet.upper() and ' + ' in snippet:
                patterns.append(r'[\'"]LIMIT[^\'\"]*[\'"]\s*\+\s*\w+')
            elif '--' in snippet and ' + ' in snippet:
                patterns.append(r'\w+\s*\+\s*[\'"].*--.*[\'"]')
            elif ('AND' in snippet.upper() or 'OR' in snippet.upper()) and '=' in snippet:
                patterns.append(r'\w+\s*\+\s*[\'"].*(?:AND|OR)\s+\d+\s*=\s*\d+')
            elif ' + ' in snippet and '"' in snippet:
                # Generic string concatenation pattern
                patterns.append(r'[\'"][^\'\"]*[\'"]\s*\+\s*\w+')
    
    return patterns + common_patterns

def highlight_sql_injection_vulnerabilities_word(code):
    """Highlight SQL injection vulnerability patterns for Word documents"""
    patterns = [
        # Updated patterns to match our fixes
        r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?\{\}.*?[\'\"]\s*\.format\s*\([^)]*\))',
        r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?%[sd].*?[\'\"]\s*%\s*[^;]+)',
        r'(f[\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?\{[^}]*\}.*?[\'"])',
        r'([\'\"]\s*.*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*[\'\"]\s*\+\s*\w+)',
        r'(\w+\s*\+\s*[\'"].*(?:SELECT|INSERT|UPDATE|DELETE|WHERE|FROM|JOIN|UNION).*[\'"])',
        r'(cursor\.execute\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(\.execute\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(cursor\.execute\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'(\.raw\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(\.raw\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'(text\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(text\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'([\'"]WHERE\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
        r'([\'"]LIKE\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
        r'([\'"]ORDER\s+BY\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(request\.args\.get\([^)]*\))',
        r'(request\.form\.get\([^)]*\))',
        r'(request\.(?:form|args)\[[^]]*\])',
        # Multi-line concatenation patterns
        r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE)[^\'\"]*[\'\"]\s*\+[^+]*\+[^+]*\+)',
        r'([\'\"]\s*WHERE[^\'\"]*[\'\"]\s*\+[^+]*\+)',
        r'([\'\"]\s*FROM[^\'\"]*[\'\"]\s*\+[^+]*\+)',
        # NoSQL injection patterns
        r'(collection\.find\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input)[^}]*\})',
        r'(\.find\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input)[^}]*\})',
        r'(db\.eval\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'(db\.eval\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(db\.eval\s*\([^)]*(?:request\.|username|user_input|\w+_input))',
        r'(\.(?:find_one|update|delete|remove|insert)\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input))',
        r'(\{[^}]*[\'"]:\s*request\.(?:form|args|json)\[[^]]*\][^}]*\})',
        r'(aggregate\s*\(\s*\[[^]]*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input))'
    ]
    
    highlighted = code
    for pattern in patterns:
        highlighted = re.sub(pattern, lambda m: f'[SQL-INJECTION-VULNERABLE:{m.group(0)}]', highlighted, flags=re.IGNORECASE | re.DOTALL)
    
    return highlighted

def scan_code_content_for_sql_injection(code_content: str, source_name: str) -> dict:
    """Scan code content for SQL injection vulnerabilities"""
    try:
        # Create temporary file for scanning
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False)
        temp_file.write(code_content)
        temp_file.close()
        
        # Create SQL injection detector and scan
        detector = SQLInjectionDetector()
        vulnerabilities = detector.scan_file(temp_file.name)
        
        # Clean up temporary file
        os.unlink(temp_file.name)
        
        # Generate highlighted code
        highlighted_code = None
        original_code = code_content
        file_name = source_name
        
        if vulnerabilities:
            highlighted_code = highlight_sql_injection_vulnerabilities(code_content, vulnerabilities)
            original_code = code_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            if '/' in source_name:
                file_name = source_name.split('/')[-1]
            elif source_name.startswith('http'):
                file_name = source_name.split('/')[-1] if '/' in source_name else 'scanned_code.py'
        
        # Calculate summary
        summary = {
            'total_vulnerabilities': len(vulnerabilities),
            'high_severity': sum(1 for v in vulnerabilities if v.severity == 'high'),
            'medium_severity': sum(1 for v in vulnerabilities if v.severity == 'medium'),
            'low_severity': sum(1 for v in vulnerabilities if v.severity == 'low'),
            'high': sum(1 for v in vulnerabilities if v.severity == 'high'),
            'medium': sum(1 for v in vulnerabilities if v.severity == 'medium'),
            'low': sum(1 for v in vulnerabilities if v.severity == 'low')
        }
        
        # Format results
        results = {
            'source': source_name,
            'scan_type': 'sql_injection',
            'summary': summary,
            'vulnerabilities': [vuln.to_dict() for vuln in vulnerabilities],
            'total_vulnerabilities': len(vulnerabilities),
            'scan_timestamp': datetime.now().isoformat(),
            'total_issues': len(vulnerabilities),
            'high_severity': summary['high_severity'],
            'medium_severity': summary['medium_severity'],
            'low_severity': summary['low_severity'],
            'high_count': summary['high'],
            'medium_count': summary['medium'],
            'low_count': summary['low'],
            'highlighted_code': highlighted_code,
            'original_code': original_code,
            'file_name': file_name
        }
        
        return results
        
    except Exception as e:
        return {
            'error': f'Error during SQL injection scan: {str(e)}',
            'source': source_name,
            'scan_type': 'sql_injection',
            'vulnerabilities': [],
            'total_vulnerabilities': 0,
            'total_issues': 0,
            'high_severity': 0,
            'medium_severity': 0,
            'low_severity': 0,
            'high_count': 0,
            'medium_count': 0,
            'low_count': 0,
            'highlighted_code': None,
            'original_code': '',
            'file_name': source_name
        }

def is_github_py_url(url):
    """Check if URL is a GitHub Python file URL"""
    return 'github.com' in url and url.endswith('.py')

def github_raw_url(url):
    """Convert GitHub blob URL to raw URL"""
    return url.replace('github.com', 'raw.githubusercontent.com').replace('/blob/', '/')

def api_scan_sql_injection(current_user):
    """API endpoint for SQL injection scanning"""
    try:
        data = request.get_json()
        code_content = data.get('code')
        url = data.get('url')
        
        if code_content:
            # Scan provided code content
            results = scan_code_content_for_sql_injection(code_content, 'Direct input')
            
        elif url:
            # Scan URL content
            if is_github_py_url(url):
                raw_url = github_raw_url(url)
                try:
                    response = requests.get(raw_url, timeout=10)
                    if response.status_code == 200:
                        results = scan_code_content_for_sql_injection(response.text, url)
                    else:
                        return jsonify({'error': f'Failed to fetch URL: {response.status_code}'}), 400
                except Exception as e:
                    return jsonify({'error': f'Error fetching URL: {str(e)}'}), 400
            else:
                return jsonify({'error': 'Invalid GitHub Python file URL'}), 400
        else:
            return jsonify({'error': 'Invalid scan parameters'}), 400
        
        return jsonify(results)
        
    except Exception as e:
        return jsonify({'error': f'Error during SQL injection scan: {str(e)}'}), 500

def api_generate_sql_injection_report(current_user):
    """API endpoint for generating SQL injection reports"""
    try:
        data = request.get_json()
        vulnerabilities = data.get('vulnerabilities', [])
        source = data.get('source', 'Unknown')
        
        if not vulnerabilities:
            return jsonify({'error': 'No vulnerabilities provided'}), 400
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('SQL Injection Security Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_heading('Report Information', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        info_table.cell(0, 0).text = 'Source'
        info_table.cell(0, 1).text = source
        info_table.cell(1, 0).text = 'Report Generated'
        info_table.cell(1, 1).text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        info_table.cell(2, 0).text = 'Total Vulnerabilities'
        info_table.cell(2, 1).text = str(len(vulnerabilities))
        info_table.cell(3, 0).text = 'Risk Level'
        info_table.cell(3, 1).text = 'High' if len(vulnerabilities) > 0 else 'Low'
        
        # Add summary
        doc.add_heading('Executive Summary', level=1)
        summary_text = f"""
This report presents the results of a comprehensive SQL injection vulnerability analysis performed on the provided code. 
The analysis identified {len(vulnerabilities)} potential SQL injection vulnerabilities that could allow attackers to manipulate 
database queries and potentially gain unauthorized access to sensitive data.

SQL injection vulnerabilities can lead to data breaches, data modification, data deletion, and in some cases, complete system compromise.
These vulnerabilities should be addressed immediately to prevent potential attacks against the application's database.
        """
        doc.add_paragraph(summary_text.strip())
        
        # Add vulnerability details
        doc.add_heading('Vulnerability Details', level=1)
        
        for i, vuln in enumerate(vulnerabilities, 1):
            doc.add_heading(f'SQL Injection Vulnerability #{i}', level=2)
            
            # Vulnerability info table
            vuln_table = doc.add_table(rows=5, cols=2)
            vuln_table.style = 'Table Grid'
            
            vuln_table.cell(0, 0).text = 'Line Number'
            vuln_table.cell(0, 1).text = str(vuln.get('line_number', 'N/A'))
            vuln_table.cell(1, 0).text = 'Severity'
            vuln_table.cell(1, 1).text = vuln.get('severity', 'Medium').title()
            vuln_table.cell(2, 0).text = 'CWE References'
            vuln_table.cell(2, 1).text = ', '.join(vuln.get('cwe_references', []))
            vuln_table.cell(3, 0).text = 'OWASP References'
            vuln_table.cell(3, 1).text = ', '.join(vuln.get('owasp_references', []))
            vuln_table.cell(4, 0).text = 'Description'
            vuln_table.cell(4, 1).text = vuln.get('description', 'SQL injection vulnerability detected')
            
            # Vulnerable code
            doc.add_paragraph('Vulnerable Code:', style='Heading 3')
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(vuln.get('code_snippet', ''))
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(10)
            code_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
            
            # Remediation
            doc.add_paragraph('Remediation:', style='Heading 3')
            remediation_para = doc.add_paragraph(vuln.get('remediation', 'Use parameterized queries and proper input validation'))
            
            doc.add_paragraph()  # Add spacing
        
        # Add recommendations
        doc.add_heading('SQL Injection Prevention Recommendations', level=1)
        recommendations = """
1. Parameterized Queries: Always use parameterized queries or prepared statements instead of string concatenation.
2. Input Validation: Validate all user inputs on both client and server side.
3. Least Privilege: Use database accounts with minimal necessary privileges.
4. Stored Procedures: Use stored procedures when possible, but ensure they are also parameterized.
5. Escape Special Characters: If parameterized queries are not possible, properly escape special characters.
6. ORM Usage: Use Object-Relational Mapping (ORM) frameworks that handle parameterization automatically.
7. Database Firewalls: Implement database firewalls to detect and block SQL injection attempts.
8. Regular Security Testing: Conduct regular security testing including automated SQL injection scanning.
9. Code Reviews: Implement thorough code review processes to catch SQL injection vulnerabilities.
10. Security Training: Provide security awareness training to developers about SQL injection risks.
        """
        doc.add_paragraph(recommendations.strip())
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'sql_injection_security_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': f'Error generating SQL injection report: {str(e)}'}), 500

def api_sql_injection_sonarqube_export(current_user):
    """API endpoint for SQL injection SonarQube export"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        vulnerabilities = data.get('vulnerabilities', [])
        if not vulnerabilities:
            return jsonify({'error': 'No vulnerabilities provided'}), 400
        
        # Convert to SonarQube format
        sonar_issues = []
        for vuln in vulnerabilities:
            # Map severity to SonarQube severity levels
            severity_map = {
                'high': 'CRITICAL',
                'medium': 'MAJOR',
                'low': 'MINOR'
            }
            sonar_severity = severity_map.get(vuln.get('severity', 'medium'), 'MAJOR')
            
            sonar_issue = {
                "engineId": "python-sql-injection-scanner",
                "ruleId": vuln.get('rule_key', 'python:S2077'),
                "severity": sonar_severity,
                "type": "VULNERABILITY",
                "primaryLocation": {
                    "message": vuln.get('description', 'SQL injection vulnerability'),
                    "filePath": vuln.get('file_path', 'unknown'),
                    "textRange": {
                        "startLine": vuln.get('line_number', 1),
                        "endLine": vuln.get('line_number', 1)
                    }
                },
                "cwe": vuln.get('cwe_references', []),
                "owasp": vuln.get('owasp_references', []),
                "confidence": vuln.get('confidence', 0.5)
            }
            sonar_issues.append(sonar_issue)
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.json')
        json.dump({"issues": sonar_issues}, temp_file, indent=2)
        temp_file.close()
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'sql_injection_sonarqube_issues_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json',
            mimetype='application/json'
        )
        
    except Exception as e:
        return jsonify({'error': f'Error exporting SQL injection SonarQube format: {str(e)}'}), 500


def test_sql_injection_scanner():
    """Test function to verify the SQL injection scanner is working"""
    test_cases = [
        # SQLAlchemy text() cases
        'query = text(f"SELECT * FROM users WHERE id = {user_id}")',
        'result = text(f"SELECT * FROM products WHERE category = {category}")',
        'sql = text("SELECT * FROM users WHERE id = " + str(user_id))',
        
        # F-string cases
        'sql = f"SELECT * FROM users WHERE name = {username}"',
        'query = f"INSERT INTO logs VALUES ({user_id}, {message})"',
        
        # String concatenation cases
        'query = "SELECT * FROM users WHERE id = " + user_id',
        'sql = "DELETE FROM users WHERE name = " + username',
        
        # cursor.execute cases
        'cursor.execute(f"SELECT * FROM users WHERE id = {user_id}")',
        'cursor.execute("SELECT * FROM users WHERE id = " + str(user_id))',
        
        # Django ORM raw cases
        'User.objects.raw(f"SELECT * FROM users WHERE id = {user_id}")',
        'results = MyModel.objects.raw("SELECT * FROM table WHERE id = " + user_id)',
        
        # NoSQL injection cases
        'collection.find({"name": request.form["username"]})',
        'db.eval(f"this.name == {username}")',
        
        # Safe cases (should NOT be detected)
        'query = text("SELECT * FROM users WHERE id = :user_id", user_id=user_id)',
        'cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))',
        'safe_query = "SELECT * FROM users"',
    ]
    
    print("Testing SQL Injection Scanner...")
    print("=" * 50)
    
    for i, test_case in enumerate(test_cases, 1):
        print(f"\nTest {i}: {test_case}")
        
        # Create temporary file with test case
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False)
        temp_file.write(test_case)
        temp_file.close()
        
        # Scan the test case
        detector = SQLInjectionDetector()
        vulnerabilities = detector.scan_file(temp_file.name)
        
        # Clean up
        os.unlink(temp_file.name)
        
        # Report results
        if vulnerabilities:
            print(f"  ✓ DETECTED: {len(vulnerabilities)} vulnerability(ies)")
            for vuln in vulnerabilities:
                print(f"    - {vuln.description}")
        else:
            # Check if this is a safe case that should NOT be detected
            if any(safe_pattern in test_case for safe_pattern in [':user_id', '%s', '(user_id,)', 'SELECT * FROM users"']):
                print(f"  ✓ SAFE: No vulnerabilities detected (expected)")
            else:
                print(f"  ✗ MISSED: No vulnerabilities detected (unexpected)")
    
    print("\n" + "=" * 50)
    print("Test completed. If you see any 'MISSED' results, there may be an issue with the scanner.")
    
    
# Test the scanner if run directly
if __name__ == "__main__":
    test_sql_injection_scanner() 


# cd /Users/mohammen.almasi/thesis/06.27 && source venv/bin/activate && cd backend && python main.py

# cd thesis/06.27/frontend && npm start